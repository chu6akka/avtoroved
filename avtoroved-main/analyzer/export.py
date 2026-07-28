"""
Модуль экспорта отчётов в DOCX и загрузки файлов.
"""
from __future__ import annotations
import os
import re
from datetime import datetime
from typing import List

from analyzer.stanza_backend import TokenInfo


def _docx_block_texts(doc):
    """
    Тексты блоков DOCX в порядке их следования в теле документа:
    обычные абзацы (w:p) и таблицы (w:tbl) — документы, свёрстанные
    таблицей (резюме, анкеты), иначе дают пустой текст.
    """
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    for child in doc.element.body.iterchildren():
        if child.tag == qn('w:p'):
            text = Paragraph(child, doc).text
            if text.strip():
                yield text
        elif child.tag == qn('w:tbl'):
            for row in Table(child, doc).rows:
                cells, seen = [], set()
                for cell in row.cells:
                    # Объединённые ячейки повторяются в row.cells — берём один раз.
                    if id(cell._tc) in seen:
                        continue
                    seen.add(id(cell._tc))
                    text = cell.text.strip()
                    if text:
                        cells.append(text)
                if cells:
                    yield '\t'.join(cells)


def load_text_from_file(filepath: str) -> str:
    """Загрузить текст из .txt или .docx файла."""
    ext = os.path.splitext(filepath)[1].lower()
    if ext == '.txt':
        for enc in ('utf-8', 'cp1251', 'cp866', 'latin-1'):
            try:
                with open(filepath, 'r', encoding=enc) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        raise ValueError(f"Не удалось прочитать: {filepath}")
    elif ext == '.docx':
        from docx import Document
        doc = Document(filepath)
        return '\n'.join(_docx_block_texts(doc))
    else:
        raise ValueError(f"Неподдерживаемый формат: {ext}")


def export_report_docx(filepath: str, text: str, metrics: dict,
                       error_result, tokens: List[TokenInfo],
                       strat_result=None, gigacheck_result=None,
                       thematic_result=None,
                       ogorelkov_result=None, ogorelkov_detailed=False):
    """Экспорт полного отчёта в DOCX."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    h = doc.add_heading('СВОДНЫЙ ОТЧЁТ АВТОРОВЕДЧЕСКОГО АНАЛИЗА', level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'Дата: {datetime.now().strftime("%d.%m.%Y %H:%M")}')

    # 1. Описание
    doc.add_heading('1. Описание речевого продукта', level=2)
    words_list = re.findall(r"[А-Яа-яЁё]+", text)
    first_w = " ".join(words_list[:6]) + "..." if len(words_list) > 6 else " ".join(words_list)
    last_w = "..." + " ".join(words_list[-6:]) if len(words_list) > 6 else ""
    doc.add_paragraph(f'Текст на русском языке. Начало: «{first_w}»')
    if last_w:
        doc.add_paragraph(f'Окончание: «{last_w}»')
    doc.add_paragraph(f'Вербальный объём: {metrics["дополнительно"]["Всего слов"]} слов.')

    # 2. Количественные характеристики
    doc.add_heading('2. Количественные характеристики', level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    table.rows[0].cells[0].text = 'Показатель'
    table.rows[0].cells[1].text = 'Значение'
    for k, v in metrics["дополнительно"].items():
        row = table.add_row().cells
        row[0].text = str(k)
        row[1].text = str(v)

    # 3. Распределение частей речи
    doc.add_heading('3. Распределение частей речи', level=2)
    if metrics["частоты"]:
        t2 = doc.add_table(rows=1, cols=3)
        t2.style = 'Light Grid Accent 1'
        t2.rows[0].cells[0].text = 'Часть речи'
        t2.rows[0].cells[1].text = 'Количество'
        t2.rows[0].cells[2].text = 'Доля'
        for p, v in metrics["частоты"].items():
            row = t2.add_row().cells
            row[0].text = p
            row[1].text = str(v['количество'])
            row[2].text = f"{v['коэффициент']:.2%}"

    # 3.1. Морфологические коэффициенты (методика САЭ)
    sae = metrics.get("sae_coefficients", {})
    sae_rows = sae.get("rows", [])
    base_cnt = sae.get("base_counts", {})
    if sae_rows:
        doc.add_heading('3.1. Морфологические коэффициенты (САЭ)', level=2)
        doc.add_paragraph(
            'По методике С.М. Вул, Е.И. Галяшиной. '
            'Местоимения = PRON + DET (соответствует традиционной русской грамматике). '
            'Глаголы — личные формы (без причастий и деепричастий).'
        )
        # Базовые счётчики
        doc.add_paragraph('Базовые счётчики:').bold = True
        t_base = doc.add_table(rows=1, cols=2)
        t_base.style = 'Light Grid Accent 1'
        t_base.rows[0].cells[0].text = 'Категория'
        t_base.rows[0].cells[1].text = 'Количество'
        for lbl, cnt in base_cnt.items():
            row = t_base.add_row().cells
            row[0].text = str(lbl)
            row[1].text = str(cnt)
        doc.add_paragraph('')
        # 20 коэффициентов
        t_sae = doc.add_table(rows=1, cols=4)
        t_sae.style = 'Light Grid Accent 1'
        hdr = t_sae.rows[0].cells
        hdr[0].text = '№'
        hdr[1].text = 'Показатель'
        hdr[2].text = 'Числ./знаменатель'
        hdr[3].text = 'Коэффициент'
        for r in sae_rows:
            row = t_sae.add_row().cells
            row[0].text = str(r["n"])
            row[1].text = r["label"]
            row[2].text = f"{r['numerator']}/{r['denominator']}"
            row[3].text = f"{r['value']:.3f}" if r["value"] is not None else "н/д"

    # 3.2. POS-биграммы
    pos_bg = metrics.get("pos_bigrams", {})
    top_bg = pos_bg.get("top_bigrams", [])
    if top_bg:
        doc.add_heading('3.2. Коэффициенты сочетаемости частеречных пар (POS-биграммы)', level=2)
        doc.add_paragraph(
            'Метод POS-биграмм: частотное распределение последовательных пар '
            'частей речи отражает грамматические привычки автора. '
            '(Litvinova et al., 2015–2016).'
        )
        t3 = doc.add_table(rows=1, cols=4)
        t3.style = 'Light Grid Accent 1'
        hdr = t3.rows[0].cells
        hdr[0].text = '№'
        hdr[1].text = 'Пара'
        hdr[2].text = 'Количество'
        hdr[3].text = 'Коэффициент'
        for i, bg in enumerate(top_bg[:15], 1):
            row = t3.add_row().cells
            row[0].text = str(i)
            row[1].text = bg["pair_full"]
            row[2].text = str(bg["count"])
            row[3].text = f'{bg["freq"]:.4f}'

    # 4. Навыки
    if error_result and error_result.skill_levels:
        doc.add_heading('4. Степени развития языковых навыков (по С.М. Вул)', level=2)
        for skill in error_result.skill_levels:
            p = doc.add_paragraph()
            run = p.add_run(f'{skill.skill_name}: ')
            run.bold = True
            p.add_run(f'{skill.level.upper()} ({skill.description})')

    # 5. Ошибки
    if error_result and error_result.errors:
        doc.add_heading('5. Выявленные речевые ошибки', level=2)
        by_type = {}
        for e in error_result.errors:
            by_type.setdefault(e.error_type, []).append(e)
        for etype, errs in by_type.items():
            doc.add_heading(f'{etype.upper()} ({len(errs)})', level=3)
            for i, e in enumerate(errs[:10], 1):
                txt = f'{i}. «{e.fragment}» — {e.description}'
                if e.suggestion:
                    txt += f' → {e.suggestion}'
                if e.rule_ref:
                    txt += f' [{e.rule_ref}]'
                doc.add_paragraph(txt)

    # 6. Лексическая стратификация
    if strat_result:
        from analyzer.stratification_engine import LAYER_META
        marked = len(strat_result.tokens)
        doc.add_heading('6. Лексическая стратификация', level=2)
        doc.add_paragraph(
            f'Маркированных единиц: {marked} из {strat_result.total_words} слов '
            f'({strat_result.marked_ratio:.1%}).')
        t4 = doc.add_table(rows=1, cols=3)
        t4.style = 'Light Grid Accent 1'
        t4.rows[0].cells[0].text = 'Пласт'
        t4.rows[0].cells[1].text = 'Кол-во'
        t4.rows[0].cells[2].text = 'Доля'
        total_w = strat_result.total_words or 1
        for layer_key, meta in sorted(
            LAYER_META.items(), key=lambda x: -x[1]["priority"]
        ):
            cnt = strat_result.layer_counts.get(layer_key, 0)
            if cnt == 0:
                continue
            row = t4.add_row().cells
            row[0].text = meta["label"]
            row[1].text = str(cnt)
            row[2].text = f"{cnt / total_w:.1%}"

    # 7. GigaCheck (если есть)
    if gigacheck_result:
        doc.add_heading('7. Анализ ИИ-генерации (GigaCheck)', level=2)
        doc.add_paragraph(
            f'Вероятность ИИ-генерации: {gigacheck_result.get("overall_score", 0):.1%}')
        doc.add_paragraph(
            'Примечание: результат является вспомогательным инструментом и '
            'не заменяет лингвистический анализ.')

    # 8. Тематические словари (если есть)
    if thematic_result:
        doc.add_heading('8. Тематическая атрибуция', level=2)
        for domain, data in list(thematic_result.items())[:3]:
            doc.add_paragraph(
                f'{data["label"]}: {data["count"]} слов '
                f'(k={data["density"]:.4f} на 1000 слов)')

    # Частотный анализ служебной лексики (Огорелков)
    if ogorelkov_result:
        doc.add_heading('Частотный анализ служебной лексики (по И.В. Огорелкову)',
                        level=2)
        doc.add_paragraph(
            'Относительные частоты (ipm) употребления служебных '
            'лексико-грамматических классов слов; нормирование по частотному '
            'словарю О.Н. Ляшевской и С.А. Шарова.')
        doc.add_paragraph(
            f'Словоупотреблений: {ogorelkov_result["total_words"]}; '
            f'словарь маркеров sha256: {ogorelkov_result["dict_sha256"][:16]}…')

        def _na(v):
            return 'н/д' if v is None else f'{v:g}'

        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        for i, h in enumerate(('Категория', 'Использовано лемм', 'Вхождения',
                               'ipm', 'Доля, %')):
            table.rows[0].cells[i].text = h
        for cat, d in ogorelkov_result["categories"].items():
            row = table.add_row().cells
            row[0].text = cat.replace('_', ' ')
            row[1].text = f'{d["used"]} из {d["total_lemmas"]}'
            row[2].text = str(d["total_count"])
            row[3].text = _na(d["total_ipm"])
            row[4].text = _na(d["share_pct"])

        if ogorelkov_detailed:
            for cat, d in ogorelkov_result["categories"].items():
                if not d["lemmas"]:
                    continue
                doc.add_paragraph(cat.replace('_', ' '), style='Heading 3')
                dt = doc.add_table(rows=1, cols=5)
                dt.style = 'Table Grid'
                for i, h in enumerate(('Лемма', 'Вхождения', 'ipm текста',
                                       'ipm НКРЯ', 'Коэф. отклонения')):
                    dt.rows[0].cells[i].text = h
                for lem, ld in d["lemmas"].items():
                    row = dt.add_row().cells
                    row[0].text = lem
                    row[1].text = str(ld["count"])
                    row[2].text = _na(ld["ipm_text"])
                    row[3].text = _na(ld["ipm_rnc"])
                    row[4].text = _na(ld["ratio"])

    # Вывод
    doc.add_heading('Вывод', level=2)
    wc = metrics["дополнительно"]["Всего слов"]
    if wc >= 500:
        doc.add_paragraph(f'Объём текста ({wc} слов) пригоден для судебной автороведческой экспертизы.')
    elif wc >= 200:
        doc.add_paragraph(f'Объём текста ({wc} слов) пригоден для предварительного анализа.')
    else:
        doc.add_paragraph(f'Объём текста ({wc} слов) недостаточен (минимум 500 слов).')
    doc.save(filepath)


def export_comparison_docx(filepath: str, structured, comp: dict,
                           text1: str, text2: str, expert_verdict: str = ""):
    """
    Экспорт сравнительного исследования в DOCX по структуре методики
    Рубцовой 2007 (ЭКЦ МВД): два комплекса признаков (совпадающие /
    различающиеся) по уровням НН/НС/НСВ, вспомогательные метрики отдельно,
    подсказка по шкале (с. 85) и поле вывода эксперта.

    Модуль не печатает решение — окончательный вывод формулирует эксперт.
    """
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    h = doc.add_heading('СРАВНИТЕЛЬНОЕ ИССЛЕДОВАНИЕ ТЕКСТОВ', level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'Дата: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
    doc.add_paragraph(
        'Методика: Рубцова И.И. и др. Комплексная методика производства '
        'судебно-автороведческих экспертиз. — М.: ЭКЦ МВД России, 2007. '
        'Сравнение по трём уровням индивидуализации речемыслительного навыка '
        '(НН — набор норм; НС — набор свойств норм; НСВ — набор средств выражения).'
    )

    # Если структурированного результата нет — деградация к минимуму
    if structured is None:
        doc.add_paragraph('Структурированный результат недоступен.')
        doc.save(filepath)
        return

    # ── Счётчики ──────────────────────────────────────────────────────────
    doc.add_heading('1. Итоги сопоставления', level=2)
    ls = structured.level_summary
    doc.add_paragraph(
        f'Всего признаков: {structured.total_features} '
        f'(совпадающих {len(structured.matches)}, различающихся {len(structured.diffs)}).'
    )
    doc.add_paragraph(
        'По уровням (совпадения/различия): '
        + ', '.join(f"{lv} {ls.get(lv, {}).get('match', 0)}/{ls.get(lv, {}).get('diff', 0)}"
                    for lv in ('НН', 'НС', 'НСВ'))
    )
    doc.add_paragraph(
        f'Высокоинформативных совпадений: {structured.high_informative_matches} '
        f'из требуемых не менее {structured.threshold} (с. 85). '
        f'Высокоинформативных различий: {structured.high_informative_diffs}.'
    )

    def _complex_table(title, feats):
        doc.add_heading(title, level=2)
        if not feats:
            doc.add_paragraph('— не выявлено —')
            return
        order = {'НН': 0, 'НС': 1, 'НСВ': 2}
        feats = sorted(feats, key=lambda f: order.get(f.level, 9))
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'
        hdr = table.rows[0].cells
        for i, t in enumerate(['Ур.', 'Признак', 'Текст 1', 'Текст 2', 'Высокоинф.']):
            hdr[i].text = t
        for f in feats:
            row = table.add_row().cells
            row[0].text = f.level
            row[1].text = f.name + (' (устойч.)' if f.stable else '')
            row[2].text = str(f.value1)
            row[3].text = str(f.value2)
            row[4].text = 'да' if f.high_informative else ''
            if f.note:
                note_p = row[1].add_paragraph()
                run = note_p.add_run(f.note)
                run.font.size = Pt(9)
                run.italic = True

    _complex_table('2. Комплекс СОВПАДАЮЩИХ признаков', structured.matches)
    _complex_table('3. Комплекс РАЗЛИЧАЮЩИХСЯ признаков', structured.diffs)

    # ── Вспомогательные метрики ───────────────────────────────────────────
    doc.add_heading('4. Вспомогательные объективизирующие показатели '
                    '(не являются выводом)', level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    table.rows[0].cells[0].text = 'Показатель'
    table.rows[0].cells[1].text = 'Значение'
    for label, key in [("Общее сходство (агрегат)", "overall"),
                       ("Лексическое (Jaccard)", "jaccard"),
                       ("Морфологическое (POS)", "pos_similarity"),
                       ("Синтаксическое", "syntactic_similarity"),
                       ("TTR-сходство", "ttr_similarity"),
                       ("POS-биграммное", "bigram_similarity"),
                       ("SBERT (семантическое)", "sbert_sim")]:
        if key in (comp or {}):
            row = table.add_row().cells
            row[0].text = label
            row[1].text = f"{comp[key]:.1%}"

    # ── Подсказка и вывод эксперта ────────────────────────────────────────
    doc.add_heading('5. Синтезирующая стадия', level=2)
    p = doc.add_paragraph()
    p.add_run('Подсказка по шкале (с. 85, НЕ вывод): ').bold = True
    p.add_run(structured.hint)
    if structured.hint_basis:
        doc.add_paragraph('Основание: ' + '; '.join(structured.hint_basis) + '.')
    doc.add_paragraph('Окончательный вывод формулирует эксперт.').italic = True

    doc.add_heading('Вывод эксперта', level=3)
    doc.add_paragraph(expert_verdict if expert_verdict.strip()
                      else '__________________________________________________')

    doc.save(filepath)


def export_morphology_docx(filepath: str, tokens: List[TokenInfo], morph_indices: dict):
    """
    Экспорт морфологической разметки + 20 индексов идиостиля в DOCX.
    Структура:
      1. Таблица морфологического разбора (Словоформа | Лемма | ЧР | Морф. признаки)
      2. Таблица морфологических индексов (20 позиций)
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # ── Заголовок ──────────────────────────────────────────────────────
    h = doc.add_heading('МОРФОЛОГИЧЕСКИЙ РАЗБОР ТЕКСТА', level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        f'Дата: {datetime.now().strftime("%d.%m.%Y %H:%M")}  |  '
        f'Токенов: {len(tokens)}'
    )

    # ── Раздел 1: Морфологическая таблица ─────────────────────────────
    doc.add_heading('1. Морфологический разбор', level=2)

    headers = ['Словоформа', 'Лемма', 'Часть речи', 'Морфологические признаки']
    col_widths_cm = [3.0, 3.0, 3.5, 8.5]

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Table Grid'

    # Заголовочная строка
    hdr_row = tbl.rows[0]
    for i, (h_text, w) in enumerate(zip(headers, col_widths_cm)):
        cell = hdr_row.cells[i]
        cell.width = int(Cm(w))
        p = cell.paragraphs[0]
        run = p.add_run(h_text)
        run.bold = True
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Серый фон
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'D9D9D9')
        tc_pr.append(shd)

    # Строки данных
    for tok in tokens:
        row_cells = tbl.add_row().cells
        vals = [tok.text, tok.lemma, tok.pos_label, tok.feats or '—']
        for i, (cell, val) in enumerate(zip(row_cells, vals)):
            cell.width = int(Cm(col_widths_cm[i]))
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(9)
            if i == 2:  # Часть речи — курсив
                run.italic = True

    doc.add_paragraph('')

    # ── Раздел 2: Морфологические индексы ─────────────────────────────
    doc.add_heading('2. Морфологические индексы идиостиля', level=2)
    doc.add_paragraph(
        'Источник: Лабораторная работа № 11, судебная автороведческая экспертиза / '
        'Соколова Т.П. (МГЮА). Индекс #7 (абстрактные/конкретные существительные) '
        'требует семантической разметки и в автоматическом режиме недоступен.'
    ).runs[0].font.size = Pt(10)

    total_w = morph_indices.get("total_words", 0)
    sent_c  = morph_indices.get("sent_count", 0)
    doc.add_paragraph(
        f'Вербальный объём: {total_w} слов  |  Предложений: {sent_c}'
    ).runs[0].font.size = Pt(10)

    idx_headers = ['№', 'Индекс', 'Числитель', 'Знаменатель', 'Значение']
    idx_widths  = [0.8, 8.2, 2.0, 2.2, 2.2]

    tbl2 = doc.add_table(rows=1, cols=5)
    tbl2.style = 'Table Grid'

    hdr2 = tbl2.rows[0]
    for i, (h_text, w) in enumerate(zip(idx_headers, idx_widths)):
        cell = hdr2.cells[i]
        cell.width = int(Cm(w))
        p = cell.paragraphs[0]
        run = p.add_run(h_text)
        run.bold = True
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'D9D9D9')
        tc_pr.append(shd)

    indices = morph_indices.get("indices", [])
    for idx, (name, num, den, val) in enumerate(indices, 1):
        row_cells = tbl2.add_row().cells
        num_s = str(num) if num is not None else '—'
        den_s = str(den) if den is not None else '—'
        val_s = str(val) if val is not None else 'нет данных'
        widths = idx_widths
        for i, (cell, txt) in enumerate(zip(row_cells,
                                            [str(idx), name, num_s, den_s, val_s])):
            cell.width = int(Cm(widths[i]))
            p = cell.paragraphs[0]
            run = p.add_run(txt)
            run.font.size = Pt(10)
            if i == 4 and val is not None:
                run.bold = True
            if i == 4 and val is None:
                run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            if i in (2, 3, 4):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(filepath)
