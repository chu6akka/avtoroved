from __future__ import annotations

import csv
import hashlib
import io
import json
from pathlib import Path
import zipfile

from .models import ExpertCase, ExpertStatus


class ReportService:
    @staticmethod
    def _evidentiary_observations(case: ExpertCase):
        for object_id, items in case.observations.items():
            for obs in items:
                if obs.eligible_for_synthesis:
                    yield object_id, obs

    def export_docx(self, case: ExpertCase, path: str | Path) -> None:
        from docx import Document
        from docx.shared import Pt

        doc = Document()
        doc.styles["Normal"].font.name = "Times New Roman"
        doc.styles["Normal"].font.size = Pt(12)
        doc.add_heading("ПРОЕКТ ЗАКЛЮЧЕНИЯ ЭКСПЕРТА-АВТОРОВЕДА", 0)
        doc.add_paragraph(f"Дело: {case.title}\nПрофиль методики: {case.method_profile}\nИдентификатор: {case.id}")
        doc.add_heading("1. Объекты исследования", level=1)
        for obj in case.objects:
            doc.add_paragraph(
                f"{obj.id}: {obj.title}; роль: {obj.role.value}; источник: {obj.source_name or 'не указан'}; "
                f"SHA-256: {obj.source_sha256 or 'не зафиксирован'}"
            )
        doc.add_heading("2. Оценка пригодности", level=1)
        if not case.suitability:
            doc.add_paragraph("Не выполнена.")
        else:
            doc.add_paragraph("Материалы пригодны." if case.suitability.suitable else "Материалы непригодны для идентификационного вывода.")
            for issue in case.suitability.issues:
                doc.add_paragraph(issue.message, style="List Bullet")
        doc.add_heading("3. Подтверждённые признаки", level=1)
        rows = list(self._evidentiary_observations(case))
        if rows:
            table = doc.add_table(rows=1, cols=5)
            table.style = "Table Grid"
            for cell, title in zip(table.rows[0].cells, ("Объект", "Признак", "Значение", "Источник", "Ограничения")):
                cell.text = title
            for object_id, obs in rows:
                cells = table.add_row().cells
                cells[0].text = object_id; cells[1].text = obs.feature_id
                cells[2].text = f"{obs.value} {obs.unit}".strip(); cells[3].text = obs.source
                cells[4].text = "; ".join(obs.limitations) or "—"
        else:
            doc.add_paragraph("Экспертом не подтверждены признаки, допустимые для синтеза.")
        doc.add_heading("4. Сравнительное исследование", level=1)
        for comparison in case.comparisons:
            doc.add_heading(f"Спорный объект {comparison.disputed_object_id}", level=2)
            for feature in comparison.features:
                doc.add_paragraph(
                    f"{feature.feature_id}: {feature.outcome}; спорный={feature.disputed_value}; "
                    f"образцы={feature.sample_mean}; квалификация="
                    f"{feature.qualification.value if feature.qualification else 'не дана'}; {feature.explanation}"
                )
        doc.add_heading("5. Вывод", level=1)
        if case.decision:
            doc.add_paragraph(f"{case.decision.verdict.value}: {case.decision.rationale}")
            doc.add_paragraph(f"Эксперт: {case.decision.expert_name}; дата: {case.decision.decided_at}")
        else:
            doc.add_paragraph("Вывод экспертом не утверждён.")
        doc.add_paragraph("Неподтверждённые машинные и LLM-подсказки в доказательную часть не включены.")
        doc.save(path)

    def export_verification_package(self, case: ExpertCase, path: str | Path, include_source_texts: bool = False) -> None:
        manifest = {
            "format": "aved-verification-1", "case_id": case.id, "title": case.title,
            "method_profile": case.method_profile,
            "objects": [{"id": o.id, "title": o.title, "role": o.role.value, "sha256": o.source_sha256} for o in case.objects],
            "decision": case.to_dict().get("decision"),
            "source_texts_included": include_source_texts,
        }
        feature_buf = io.StringIO()
        writer = csv.writer(feature_buf)
        writer.writerow(["object_id", "feature_id", "value", "unit", "expert_status", "source", "limitations"])
        for oid, observations in case.observations.items():
            for o in observations:
                writer.writerow([oid, o.feature_id, o.value, o.unit, o.expert_status.value, o.source, "; ".join(o.limitations)])
        files = {
            "manifest.json": json.dumps(manifest, ensure_ascii=False, indent=2).encode("utf-8"),
            "audit.json": json.dumps([e.__dict__ for e in case.audit], ensure_ascii=False, indent=2).encode("utf-8"),
            "features.csv": feature_buf.getvalue().encode("utf-8-sig"),
        }
        if include_source_texts:
            for obj in case.objects:
                files[f"texts/{obj.id}.txt"] = obj.text.encode("utf-8")
        checksums = "".join(f"{hashlib.sha256(data).hexdigest()}  {name}\n" for name, data in sorted(files.items()))
        files["SHA256SUMS.txt"] = checksums.encode("ascii")
        with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zf:
            for name, data in files.items():
                zf.writestr(name, data)
