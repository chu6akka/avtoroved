"""
Диалог пакетной обработки нескольких файлов.
"""
from __future__ import annotations
import os
from typing import List
from PyQt6.QtWidgets import (
    QDialog, QVBoxLayout, QHBoxLayout, QListWidget,
    QListWidgetItem, QPushButton, QProgressBar, QLabel,
    QFileDialog, QTextEdit, QMessageBox, QGroupBox,
    QComboBox, QSizePolicy
)
from PyQt6.QtCore import Qt, QThread, pyqtSignal
from PyQt6.QtGui import QColor


class BatchWorkerThread(QThread):
    """Поток для обработки файлов."""
    progress = pyqtSignal(int, int, str)   # current, total, filename
    file_done = pyqtSignal(str, dict)       # filepath, result_summary
    all_done = pyqtSignal(list)             # list of results
    error = pyqtSignal(str, str)            # filepath, error_msg

    def __init__(self, filepaths: List[str], stanza_backend, mode: str):
        super().__init__()
        self.filepaths = filepaths
        self.stanza = stanza_backend
        self.mode = mode  # "errors_only" | "full" | "stats_only"

    def run(self):
        from analyzer.export import load_text_from_file
        from analyzer.metrics import calculate_metrics
        from analyzer.errors import ErrorAnalyzer

        error_analyzer = ErrorAnalyzer()
        results = []

        for i, fp in enumerate(self.filepaths):
            fname = os.path.basename(fp)
            self.progress.emit(i + 1, len(self.filepaths), fname)
            try:
                text = load_text_from_file(fp)
                tokens = self.stanza.analyze(text)
                metrics = calculate_metrics(tokens, text)
                error_result = error_analyzer.analyze(text) if self.mode != "stats_only" else None

                summary = {
                    "filepath": fp,
                    "filename": fname,
                    "text": text,
                    "tokens": tokens,
                    "metrics": metrics,
                    "error_result": error_result,
                    "word_count": metrics["дополнительно"].get("Всего слов", 0),
                    "ttr": metrics["дополнительно"].get("Лексическое разнообразие (TTR)", 0),
                    "error_count": len(error_result.errors) if error_result else 0,
                }
                results.append(summary)
                self.file_done.emit(fp, summary)
            except Exception as e:
                self.error.emit(fp, str(e))

        self.all_done.emit(results)


class BatchDialog(QDialog):
    """Диалог пакетной обработки файлов."""

    def __init__(self, stanza_backend, parent=None):
        super().__init__(parent)
        self.stanza = stanza_backend
        self._results = []
        self._worker = None
        self.setWindowTitle("Пакетная обработка файлов")
        self.setMinimumSize(700, 550)
        self._setup_ui()

    def _setup_ui(self):
        layout = QVBoxLayout(self)
        layout.setSpacing(8)

        # Список файлов
        files_group = QGroupBox("Файлы для обработки")
        files_layout = QVBoxLayout(files_group)

        self.file_list = QListWidget()
        self.file_list.setAcceptDrops(True)
        self.file_list.setSelectionMode(QListWidget.SelectionMode.ExtendedSelection)
        files_layout.addWidget(self.file_list)

        btn_row = QHBoxLayout()
        btn_add = QPushButton("➕ Добавить файлы")
        btn_add.clicked.connect(self._add_files)
        btn_remove = QPushButton("➖ Удалить выбранные")
        btn_remove.setObjectName("secondary")
        btn_remove.clicked.connect(self._remove_selected)
        btn_clear_list = QPushButton("🗑 Очистить список")
        btn_clear_list.setObjectName("secondary")
        btn_clear_list.clicked.connect(self.file_list.clear)
        btn_row.addWidget(btn_add)
        btn_row.addWidget(btn_remove)
        btn_row.addWidget(btn_clear_list)
        btn_row.addStretch()
        files_layout.addLayout(btn_row)
        layout.addWidget(files_group)

        # Режим анализа
        mode_group = QGroupBox("Режим анализа")
        mode_layout = QHBoxLayout(mode_group)
        mode_layout.addWidget(QLabel("Тип анализа:"))
        self.mode_combo = QComboBox()
        self.mode_combo.addItem("Полный анализ (метрики + ошибки)", "full")
        self.mode_combo.addItem("Только ошибки и навыки", "errors_only")
        self.mode_combo.addItem("Только статистика (быстро)", "stats_only")
        mode_layout.addWidget(self.mode_combo)
        mode_layout.addStretch()
        layout.addWidget(mode_group)

        # Прогресс
        self.progress_bar = QProgressBar()
        self.progress_bar.setRange(0, 100)
        self.progress_label = QLabel("Готов к обработке")
        layout.addWidget(self.progress_bar)
        layout.addWidget(self.progress_label)

        # Результаты
        result_group = QGroupBox("Результаты")
        result_layout = QVBoxLayout(result_group)
        self.result_text = QTextEdit()
        self.result_text.setReadOnly(True)
        self.result_text.setMaximumHeight(150)
        result_layout.addWidget(self.result_text)
        layout.addWidget(result_group)

        # Кнопки
        bottom_row = QHBoxLayout()
        self.btn_run = QPushButton("▶ Запустить обработку")
        self.btn_run.clicked.connect(self._run_batch)
        self.btn_export = QPushButton("📄 Экспорт отчёта в DOCX")
        self.btn_export.setObjectName("secondary")
        self.btn_export.setEnabled(False)
        self.btn_export.clicked.connect(self._export_docx)
        btn_close = QPushButton("Закрыть")
        btn_close.setObjectName("secondary")
        btn_close.clicked.connect(self.close)
        bottom_row.addWidget(self.btn_run)
        bottom_row.addWidget(self.btn_export)
        bottom_row.addStretch()
        bottom_row.addWidget(btn_close)
        layout.addLayout(bottom_row)

    def _add_files(self):
        fps, _ = QFileDialog.getOpenFileNames(
            self, "Выберите файлы", "",
            "Тексты (*.txt *.docx);;Все файлы (*)")
        for fp in fps:
            existing = [self.file_list.item(i).data(Qt.ItemDataRole.UserRole)
                        for i in range(self.file_list.count())]
            if fp not in existing:
                item = QListWidgetItem(os.path.basename(fp))
                item.setData(Qt.ItemDataRole.UserRole, fp)
                self.file_list.addItem(item)

    def _remove_selected(self):
        for item in self.file_list.selectedItems():
            self.file_list.takeItem(self.file_list.row(item))

    def _run_batch(self):
        if self.file_list.count() == 0:
            QMessageBox.warning(self, "Нет файлов", "Добавьте файлы для обработки.")
            return

        filepaths = [self.file_list.item(i).data(Qt.ItemDataRole.UserRole)
                     for i in range(self.file_list.count())]
        mode = self.mode_combo.currentData()

        self.btn_run.setEnabled(False)
        self.btn_export.setEnabled(False)
        self.progress_bar.setValue(0)
        self.result_text.clear()
        self._results = []

        self._worker = BatchWorkerThread(filepaths, self.stanza, mode)
        self._worker.progress.connect(self._on_progress)
        self._worker.file_done.connect(self._on_file_done)
        self._worker.all_done.connect(self._on_all_done)
        self._worker.error.connect(self._on_error)
        self._worker.start()

    def _on_progress(self, current: int, total: int, fname: str):
        pct = int(current / total * 100)
        self.progress_bar.setValue(pct)
        self.progress_label.setText(f"Обрабатывается {current}/{total}: {fname}")

    def _on_file_done(self, fp: str, summary: dict):
        self._results.append(summary)
        line = (f"✓ {summary['filename']}: "
                f"{summary['word_count']} слов, "
                f"TTR={summary['ttr']:.3f}, "
                f"ошибок: {summary['error_count']}")
        self.result_text.append(line)

    def _on_all_done(self, results: list):
        self._results = results
        self.btn_run.setEnabled(True)
        self.btn_export.setEnabled(len(results) > 0)
        self.progress_label.setText(f"Готово: обработано {len(results)} файлов")
        self.progress_bar.setValue(100)

    def _on_error(self, fp: str, msg: str):
        self.result_text.append(f"✗ {os.path.basename(fp)}: {msg}")

    def _export_docx(self):
        if not self._results:
            return
        fp, _ = QFileDialog.getSaveFileName(
            self, "Сохранить отчёт", "пакетный_отчёт.docx",
            "Word (*.docx)")
        if not fp:
            return
        try:
            self._write_batch_docx(fp)
            QMessageBox.information(self, "Готово", f"Отчёт сохранён:\n{fp}")
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", str(e))

    def _write_batch_docx(self, filepath: str):
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from datetime import datetime

        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)

        h = doc.add_heading('ПАКЕТНЫЙ АВТОРОВЕДЧЕСКИЙ АНАЛИЗ', level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f'Дата: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
        doc.add_paragraph(f'Файлов обработано: {len(self._results)}')

        # Сводная таблица
        doc.add_heading('Сводная таблица', level=2)
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'
        hdr = table.rows[0].cells
        hdr[0].text = 'Файл'
        hdr[1].text = 'Слов'
        hdr[2].text = 'TTR'
        hdr[3].text = 'Ошибок'
        hdr[4].text = 'Пригодность'
        for r in self._results:
            row = table.add_row().cells
            row[0].text = r["filename"]
            row[1].text = str(r["word_count"])
            row[2].text = str(r["ttr"])
            row[3].text = str(r["error_count"])
            wc = r["word_count"]
            row[4].text = "Пригоден" if wc >= 500 else "Предвар." if wc >= 200 else "Недост."

        # Подробные разделы
        for r in self._results:
            doc.add_heading(r["filename"], level=2)
            if r.get("error_result"):
                for skill in r["error_result"].skill_levels:
                    doc.add_paragraph(
                        f'{skill.skill_name}: {skill.level.upper()} — {skill.description}')
            doc.add_paragraph(
                f'Слов: {r["word_count"]}, TTR: {r["ttr"]}, Ошибок: {r["error_count"]}')

        doc.save(filepath)
