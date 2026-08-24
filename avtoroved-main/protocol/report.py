"""
protocol/report.py — экспорт заключения эксперта в DOCX (стадия 4).

Структура — по Приложению 1 методики Рубцовой 2007 (с.108–109):
вводная часть (эксперт, основание, объекты, вопросы) → ИССЛЕДОВАНИЕ
(четыре стадии: пригодность, раздельное, сравнительное, оценка) → ВЫВОДЫ.
В конце — техническая справка воспроизводимости (версии, хэши).

Экспорт требует зафиксированного вывода (conclusions). Файл регистрируется
в таблице reports с sha256 + запись в audit_log.
"""
from __future__ import annotations

import json
from typing import Optional

from protocol import db as protocol_db
from protocol import comparison as cmp
from protocol import conclusion as concl
from protocol import detector_filter
from protocol import feature_map as fm


def export_conclusion_docx(
    pdb: "protocol_db.ProtocolDB",
    project_id: int,
    doc_a: int,
    doc_b: int,
    filepath: str,
    header: Optional[dict] = None,     # {expert_name, case_number, questions}
    program_version: Optional[str] = None,
    ogorelkov_detailed: bool = False,  # включать полемную таблицу служебной лексики
) -> dict:
    """Собрать и сохранить заключение по паре. Возвращает сводку экспорта."""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    conclusion_row = pdb.fetch_conclusion(doc_a, doc_b)
    if conclusion_row is None:
        raise ValueError("Вывод по паре не зафиксирован — сначала зафиксируйте "
                         "форму вывода во вкладке «Вывод и заключение».")

    header = header or {}
    project = pdb.get_project(project_id)
    da, db_ = pdb.get_document(doc_a), pdb.get_document(doc_b)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # ── Вводная часть ────────────────────────────────────────────────────────
    h = doc.add_heading("ЗАКЛЮЧЕНИЕ ЭКСПЕРТА", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if header.get("case_number"):
        p = doc.add_paragraph(f"№ {header['case_number']}")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Проект: «{project['name']}» (создан {project['created_at']})")
    expert = header.get("expert_name") or project["expert_name"] or "—"
    doc.add_paragraph(f"Эксперт: {expert}")
    doc.add_paragraph(f"Заключение сформировано программой «Авторовед» "
                      f"v{program_version or project['program_version'] or '—'}, "
                      f"дата фиксации вывода: {conclusion_row['decided_at']}.")

    doc.add_heading("На экспертизу представлены", level=2)
    for d, role_note in ((da, "спорный текст"), (db_, "образец")):
        doc.add_paragraph(
            f"• {d['filename']} — {role_note}; происхождение: {d['provenance'] or '—'}; "
            f"жанр: {d['genre'] or '—'}; объём: {d['word_count'] or 0} словоформ; "
            f"SHA-256: {d['file_sha256']}")

    doc.add_heading("Перед экспертом поставлены вопросы", level=2)
    doc.add_paragraph(header.get("questions")
                      or "Является ли автор спорного текста и лицо, выполнившее "
                         "представленный образец, одним и тем же лицом?")

    # ── ИССЛЕДОВАНИЕ ─────────────────────────────────────────────────────────
    doc.add_heading("ИССЛЕДОВАНИЕ", level=1)

    # Стадия 1: пригодность.
    doc.add_heading("1. Оценка пригодности объектов", level=2)
    su_rows = [r for r in pdb.fetch_suitability(project_id)
               if r["document_id"] in (doc_a, doc_b)
               or (r["pair_doc_a"] == doc_a and r["pair_doc_b"] == doc_b)]
    if not su_rows:
        doc.add_paragraph("Стадия пригодности по паре не проводилась.")
    for r in su_rows:
        target = (f"документ #{r['document_id']}" if r["document_id"]
                  else "пара спорный↔образец")
        flags = json.loads(r["flags"]) if r["flags"] else []
        flag_txt = ("; ".join(f["message"] for f in flags)) or "без ограничений"
        doc.add_paragraph(f"{target}: {r['verdict']} ({flag_txt})")
    blocks = cmp.pair_blocks_strong_conclusion(pdb, project_id, doc_a, doc_b)
    if blocks:
        doc.add_paragraph(
            "На стадии оценки пригодности зафиксированы ограничения исследуемого "
            "материала. Их влияние на полноту исследования, достаточность "
            "установленной совокупности признаков и форму вывода оценивается экспертом.")

    # Стадия 2: раздельное исследование.
    doc.add_heading("2. Раздельное исследование", level=2)
    for d in (da, db_):
        accepted = [f for f in pdb.fetch_features(document_id=d["id"])
                    if f["status"] == fm.STATUS_ACCEPTED]
        cand_total = len([c for c in pdb.fetch_feature_candidates(d["id"])
                          if c["kind"] == "кандидат_признак"])
        doc.add_paragraph(
            f"{d['filename']}: профиль построен, кандидатов признаков {cand_total}, "
            f"принято экспертом {len(accepted)}.")

    # Стадия 3: сравнительное исследование — таблица подтверждённых позиций.
    doc.add_heading("3. Сравнительное исследование", level=2)
    confirmed = [r for r in pdb.fetch_comparisons(doc_a, doc_b)
                 if r["status"] == cmp.STATUS_CONFIRMED]
    if confirmed:
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, t in enumerate(("Признак", "Тип", "Уровень",
                               "Идентификационная значимость", "Примечание")):
            hdr[i].text = t
        for r in confirmed:
            row = table.add_row().cells
            row[0].text = r["label"] or ""
            row[1].text = r["match_type"]
            row[2].text = r["level"] or "—"
            row[3].text = r["identification_value"] or "без оценки"
            row[4].text = r["expert_note"] or ""
    else:
        doc.add_paragraph("Подтверждённых позиций сравнения нет.")

    # Сопоставление частот служебной лексики (Огорелков).
    og_cmp = cmp.ogorelkov_for_pair(pdb, doc_a, doc_b)
    if og_cmp:
        doc.add_heading('Сопоставление частот служебной лексики '
                        '(по И.В. Огорелкову)', level=3)
        doc.add_paragraph(
            'Относительные частоты (ipm) употребления служебных '
            'лексико-грамматических классов слов; нормирование по частотному '
            'словарю О.Н. Ляшевской и С.А. Шарова. A — спорный текст, '
            'B — образец. Приведены наблюдаемые величины; агрегированная мера '
            'сходства не вычисляется.')

        def _na(v):
            return '—' if v is None else f'{v:g}'

        cat_t = doc.add_table(rows=1, cols=7)
        cat_t.style = "Table Grid"
        for i, h in enumerate(('Категория', 'ipm A', 'ipm B', 'ipm НКРЯ',
                               'коэф. A', 'коэф. B', 'разность ipm A−B')):
            cat_t.rows[0].cells[i].text = h
        for r in og_cmp["categories"]:
            row = cat_t.add_row().cells
            row[0].text = r["category"].replace("_", " ")
            row[1].text = _na(r["ipm_a"])
            row[2].text = _na(r["ipm_b"])
            row[3].text = _na(r["ipm_rnc"])
            row[4].text = _na(r["ratio_a"])
            row[5].text = _na(r["ratio_b"])
            row[6].text = _na(r["diff_ipm"])

        if ogorelkov_detailed and og_cmp["lemmas"]:
            doc.add_paragraph('Полемное сопоставление '
                              '(сортировка по модулю разности ipm):')
            lem_t = doc.add_table(rows=1, cols=8)
            lem_t.style = "Table Grid"
            for i, h in enumerate(('Лемма', 'вхожд. A', 'ipm A', 'вхожд. B',
                                   'ipm B', 'ipm НКРЯ', 'коэф. A', 'коэф. B')):
                lem_t.rows[0].cells[i].text = h
            for r in og_cmp["lemmas"]:
                row = lem_t.add_row().cells
                row[0].text = r["lemma"]
                row[1].text = str(r["count_a"])
                row[2].text = _na(r["ipm_a"])
                row[3].text = str(r["count_b"])
                row[4].text = _na(r["ipm_b"])
                row[5].text = _na(r["ipm_rnc"])
                row[6].text = _na(r["ratio_a"])
                row[7].text = _na(r["ratio_b"])

    # Стадия 4: оценка результатов.
    doc.add_heading("4. Оценка результатов", level=2)
    bd = json.loads(conclusion_row["stats_snapshot"]) if conclusion_row["stats_snapshot"] else {}
    if "rubtsova" in bd:
        bd = bd["rubtsova"]
    if bd:
        coin, diff = bd.get("coincidence", {}), bd.get("difference", {})
        doc.add_paragraph(
            f"Подтверждено позиций: {bd.get('total_confirmed', 0)}; "
            f"совпадений {bd.get('total_coincidence', 0)} "
            f"(НН {coin.get('НН', 0)}, НС {coin.get('НС', 0)}, НСВ {coin.get('НСВ', 0)}, "
            f"без уровня {bd.get('coincidence_nolevel', 0)}); "
            f"различий {bd.get('total_difference', 0)} "
            f"(НН {diff.get('НН', 0)}, НС {diff.get('НС', 0)}, НСВ {diff.get('НСВ', 0)}, "
            f"без уровня {bd.get('difference_nolevel', 0)}). "
            "Форма вывода определена экспертом; программная рекомендация не формировалась.")
    # ── ВЫВОДЫ ───────────────────────────────────────────────────────────────
    doc.add_heading("ВЫВОДЫ", level=1)
    doc.add_paragraph(concl.FORM_LABELS.get(conclusion_row["form"],
                                            conclusion_row["form"]) + ".")
    if conclusion_row["justification"]:
        doc.add_paragraph(f"Обоснование эксперта: {conclusion_row['justification']}")

    # ── Техническая справка воспроизводимости ────────────────────────────────
    doc.add_heading("Техническая справка (воспроизводимость)", level=2)
    _cfg, cfg_hash = detector_filter.load_config()
    try:
        from analyzer.punct_checker import RULES_VERSION as _punct_ver
    except Exception:
        _punct_ver = "—"
    doc.add_paragraph(
        f"Версия программы: {program_version or '—'}; "
        f"хэш конфига фильтра детектора: {cfg_hash}; "
        f"версия правил пунктуации: {_punct_ver}. "
        f"Полный журнал действий хранится в базе протокола (audit_log).")

    doc.save(filepath)

    # Регистрация экспорта.
    from protocol.ingest import file_sha256
    sha = file_sha256(filepath)
    pdb.record_report(project_id, filepath, sha, pair_doc_a=doc_a,
                      pair_doc_b=doc_b, program_version=program_version)
    pdb.log_action(
        "экспортировано заключение", project_id=project_id,
        details={"pair_doc_a": doc_a, "pair_doc_b": doc_b,
                 "файл": filepath, "sha256": sha,
                 "форма_вывода": conclusion_row["form"]},
        program_version=program_version)
    return {"filepath": filepath, "sha256": sha, "form": conclusion_row["form"]}
