"""Стадия 4: программа проверяет условия, форму выбирает только эксперт."""
import json
import os

import pytest

from protocol import comparison as cmp
from protocol import conclusion as concl
from protocol import db as protocol_db
from tests.method_feature_helpers import qualified_feature


@pytest.fixture()
def pdb(tmp_path):
    return protocol_db.ProtocolDB(str(tmp_path / "concl.db"))


def _pair(pdb):
    pid = pdb.create_project("Дело", expert_name="Иванов И.И.")
    a = pdb.add_document(pid, "sporny.txt", protocol_db.ROLE_DISPUTED,
                         file_sha256="a", word_count=600)
    b = pdb.add_document(pid, "obrazec.txt", protocol_db.ROLE_SAMPLE,
                         file_sha256="b", word_count=600)
    return pid, a, b


def _position(pdb, pid, a, b, label, match=cmp.MATCH_COINCIDENCE,
              level="НСВ", group="языковые", subgroup="лексические",
              identification_value="высокая"):
    fa = qualified_feature(
        pdb, pid, a, label, group=group, subgroup=subgroup,
        expert_value=identification_value, suffix=f"a-{label}")
    fb = qualified_feature(
        pdb, pid, b, label, group=group, subgroup=subgroup,
        expert_value=identification_value, suffix=f"b-{label}")
    key = cmp.position_key(a, b, group, subgroup or "", label)
    pdb.replace_auto_comparisons(pid, a, b, [{
        "position_key": key, "feature_key_a": fa, "feature_key_b": fb,
        "group_name": group, "subgroup": subgroup, "label": label,
        "value_a": "v", "value_b": "v", "fragment_a": None,
        "fragment_b": None, "match_type": match,
    }])
    kwargs = {}
    if match != cmp.MATCH_COINCIDENCE:
        kwargs = {"difference_qualification": "SUBSTANTIAL",
                  "opportunity_status": "SUFFICIENT",
                  "expert_note": "мотивированное различие"}
    cmp.decide(pdb, pid, a, b, key, match_type=match, level=level,
               identification_value=identification_value, **kwargs)
    return key


def test_probable_thresholds_are_independent_values(pdb):
    assert cmp.CATEGORY_MIN_PROBABLE["текстологические"] == 2
    assert cmp.CATEGORY_MIN_PROBABLE[
        "языковые/орфографические+пунктуационные"] == 2
    pid, a, b = _pair(pdb)
    for i in range(2):
        _position(pdb, pid, a, b, f"t{i}", group="текстологические",
                  subgroup="архитектоника")
        _position(pdb, pid, a, b, f"o{i}", group="языковые",
                  subgroup="орфографические")
    checks = concl.methodological_checks(pdb, pid, a, b)
    cats = checks["moiseeva_ogorelkov"]["coincidences_by_category"]
    assert cats["текстологические"]["probable"]["condition_met"] is True
    assert cats["языковые/орфографические+пунктуационные"][
        "probable"]["condition_met"] is True


def test_positions_without_level_are_counted_separately(pdb):
    pid, a, b = _pair(pdb)
    _position(pdb, pid, a, b, "coin", level="")
    _position(pdb, pid, a, b, "diff", match=cmp.MATCH_DIFFERENCE, level="")
    rub = concl.methodological_checks(pdb, pid, a, b)["rubtsova"]
    assert rub["coincidence_nolevel"] == 1
    assert rub["difference_nolevel"] == 1
    assert rub["total_coincidence"] == 1
    assert rub["total_difference"] == 1
    assert rub["levels_with_coincidence"] == []
    assert rub["levels_with_difference"] == []


def test_twenty_low_matches_are_not_high_information(pdb):
    pid, a, b = _pair(pdb)
    for i in range(20):
        _position(pdb, pid, a, b, str(i), identification_value="низкая")
    checks = concl.methodological_checks(pdb, pid, a, b)
    data = checks["moiseeva_ogorelkov"]["by_category"]["языковые/лексические"]
    assert data["coincidence"] == 20
    assert data["high_identification_value_coincidence"] == 0
    assert cmp.stats(pdb, pid, a, b)["высокоинформативных_совпадений"] == 0


def test_vula_is_condition_not_conclusion(pdb):
    pid, a, b = _pair(pdb)
    label = "Общий признак: грамматический навык"
    key = cmp.position_key(a, b, "языковые", "грамматический", label)
    pdb.replace_auto_comparisons(pid, a, b, [{
        "position_key": key, "feature_key_a": None, "feature_key_b": None,
        "group_name": "языковые", "subgroup": "грамматический", "label": label,
        "value_a": "1 ошибок/200", "value_b": "5 ошибок/200",
        "fragment_a": None, "fragment_b": None, "match_type": cmp.GEN_HIGHER,
    }])
    checks = concl.methodological_checks(pdb, pid, a, b)
    assert checks["vula"]["condition_met"] is True
    assert "формализованного условия правила Вула" in checks["vula"]["note"]
    encoded = json.dumps(checks, ensure_ascii=False).lower()
    assert "категорический отрицательный" not in encoded
    assert "recommended_form" not in encoded


def test_difference_significance_and_observed_conditions_are_neutral(pdb):
    pid, a, b = _pair(pdb)
    for i in range(5):
        _position(pdb, pid, a, b, f"high-{i}", match=cmp.MATCH_DIFFERENCE,
                  identification_value="высокая")
    for i in range(3):
        _position(pdb, pid, a, b, f"low-{i}", match=cmp.MATCH_DIFFERENCE,
                  identification_value="низкая")
    mo = concl.methodological_checks(pdb, pid, a, b)["moiseeva_ogorelkov"]
    assert mo["difference_significance_counts"] == {
        "low": 3, "medium": 0, "high": 5, "unset": 0, "total": 8}
    assert mo["observed_conditions"]["high_differences_at_least_5"] is True
    assert mo["observed_conditions"]["low_differences_not_more_than_3"] is True
    assert "form" not in mo
    assert pdb.fetch_conclusion(a, b) is None

    _position(pdb, pid, a, b, "low-4", match=cmp.MATCH_DIFFERENCE,
              identification_value="низкая")
    observed = concl.methodological_checks(pdb, pid, a, b)[
        "moiseeva_ogorelkov"]["observed_conditions"]
    assert observed["low_differences_not_more_than_3"] is False


def test_no_differences_condition(pdb):
    pid, a, b = _pair(pdb)
    observed = concl.methodological_checks(pdb, pid, a, b)[
        "moiseeva_ogorelkov"]["observed_conditions"]
    assert observed["no_differences"] is True


def test_probable_negative_references_for_all_categories(pdb):
    specs = {
        "смысловые": ("смысловые", "тематические"),
        "текстологические": ("текстологические", "архитектоника"),
        "языковые/лексические": ("языковые", "лексические"),
        "языковые/стилистические": ("языковые", "стилистические"),
        "языковые/синтаксические": ("языковые", "синтаксические"),
        "языковые/орфографические+пунктуационные": ("языковые", "пунктуационные"),
        "психолингвистические": ("психолингвистические", None),
    }
    pid, a, b = _pair(pdb)
    for category, (group, subgroup) in specs.items():
        for i in range(cmp.CATEGORY_MIN_PROBABLE_NEGATIVE[category]):
            _position(pdb, pid, a, b, f"{category}-{i}",
                      match=cmp.MATCH_DIFFERENCE, group=group, subgroup=subgroup)
    differences = concl.methodological_checks(pdb, pid, a, b)[
        "moiseeva_ogorelkov"]["differences_by_category"]
    assert all(data["probable_negative"]["condition_met"]
               for data in differences.values())


@pytest.mark.parametrize("form", concl.FORMS)
def test_expert_can_choose_any_form_without_recommendation(pdb, form):
    pid, a, b = _pair(pdb)
    out = concl.decide(pdb, pid, a, b, form, program_version="6.0")
    assert out["form"] == form
    row = pdb.fetch_conclusion(a, b)
    assert row["recommended_form"] is None
    assert json.loads(row["stats_snapshot"])["rubtsova"]


def test_conclusion_and_comparison_history_append_only(pdb):
    pid, a, b = _pair(pdb)
    key = _position(pdb, pid, a, b, "x", identification_value="низкая")
    cmp.decide(pdb, pid, a, b, key, match_type=cmp.MATCH_COINCIDENCE,
               level="НСВ", identification_value="высокая")
    concl.decide(pdb, pid, a, b, concl.FORM_NPV)
    concl.decide(pdb, pid, a, b, concl.FORM_POS_PROBABLE, "оценка эксперта")
    assert len(pdb.fetch_comparison_decisions(a, b)) == 2
    assert len(pdb.fetch_conclusion_decisions(a, b)) == 2


def test_invalid_form(pdb):
    pid, a, b = _pair(pdb)
    with pytest.raises(ValueError):
        concl.decide(pdb, pid, a, b, "может_быть")


def test_export_creates_docx_without_program_recommendation(pdb, tmp_path):
    from protocol.report import export_conclusion_docx
    pid, a, b = _pair(pdb)
    _position(pdb, pid, a, b, "без оценки", identification_value="")
    pdb.save_suitability(pid, verdict="пригоден_с_ограничениями",
                         blocks_strong_conclusion=True, document_id=a,
                         flags=[], metrics={})
    concl.decide(pdb, pid, a, b, concl.FORM_NPV, "Недостаточность данных")
    fp = str(tmp_path / "заключение.docx")
    summary = export_conclusion_docx(pdb, pid, a, b, fp)
    assert os.path.exists(fp) and len(summary["sha256"]) == 64
    from docx import Document
    document = Document(fp)
    text = "\n".join(p.text for p in document.paragraphs)
    assert "Рекомендация по правилу" not in text
    assert "категорическая форма вывода недоступна" not in text.lower()
    assert "заблокирована программой" not in text.lower()
    table_text = [[cell.text for cell in row.cells]
                  for table in document.tables for row in table.rows]
    assert any("Идентификационная значимость" in row for row in table_text)
    assert any("без оценки" in row for row in table_text)
