#!/usr/bin/env python3
"""
Автороведческий анализатор текста v4 (Stanza)
=============================================
GUI-приложение для судебно-автороведческой экспертизы русского текста.

Бэкенд: Stanford Stanza (морфология, лемматизация, POS-теги)
Экспорт: DOCX (python-docx)
Сравнение: сравнение целых текстов с выделением совпадений

Методическая основа:
 - С.М. Вул «Судебно-автороведческая идентификационная экспертиза» (2007)
 - И.Р. Гальперин «Текст как объект лингвистического исследования» (2007)
 - Litvinova et al. (2015–2016) — POS-биграммы для профилирования автора
 - «Правила русской орфографии и пунктуации» (утв. АН СССР, 1956)

Новое в v4:
 - Коэффициенты сочетаемости частеречных пар (POS-биграммы)
 - Тепловая карта (heatmap) POS-биграмм
 - Расширенные пунктуационные проверки по Правилам 1956 г.
 - Контекстные меню (Ctrl+C/V/X, правый клик)

Запуск:
  pip install -r requirements.txt
  python autoroved_gui.py
"""

from __future__ import annotations
import colorsys, math, os, re, statistics, sys, threading
import tkinter as tk
from tkinter import messagebox, ttk, scrolledtext, filedialog
from collections import Counter
from dataclasses import dataclass
from typing import List, Dict, Tuple
from datetime import datetime

# Добавляем текущую директорию для импорта analyzer/
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from analyzer.errors import (ErrorAnalyzer, ErrorAnalysisResult, format_error_report,
                              format_internet_profile, analyze_internet_communication)

# Модуль лексической стратификации (подключается если словарь рядом с программой)
try:
    from lexical_stratification import (
        analyze_stratification, format_stratification_report,
        StratificationResult, LAYER_LABELS, LAYER_ORDER, LAYER_COLORS,
    )
    _STRATIFICATION_AVAILABLE = True
except ImportError:
    _STRATIFICATION_AVAILABLE = False

WORD_RE = re.compile(r"[A-Za-zА-Яа-яЁё]+")

# ============================================================================
# СЛОВАРИ ПЕРЕВОДА UPOS
# ============================================================================
UPOS_RU = {
    "NOUN": "Существительное", "PROPN": "Имя собственное", "VERB": "Глагол",
    "AUX": "Вспомогательный глагол", "ADJ": "Прилагательное", "ADV": "Наречие",
    "PRON": "Местоимение", "NUM": "Числительное", "DET": "Определительное слово",
    "ADP": "Предлог", "PART": "Частица", "CCONJ": "Сочинительный союз",
    "SCONJ": "Подчинительный союз", "INTJ": "Междометие", "PUNCT": "Пунктуация",
    "SYM": "Символ", "X": "Другое",
}

# Краткие POS-метки для таблиц и heatmap
UPOS_SHORT = {
    "NOUN": "СУЩ", "PROPN": "ИМЯ", "VERB": "ГЛ", "AUX": "ВСПГЛ",
    "ADJ": "ПРИЛ", "ADV": "НАР", "PRON": "МЕСТ", "NUM": "ЧИСЛ",
    "DET": "ОПР", "ADP": "ПРЕДЛ", "PART": "ЧАСТ", "CCONJ": "ССОЮЗ",
    "SCONJ": "ПСОЮЗ", "INTJ": "МЕЖД", "PUNCT": "ПУНКТ", "SYM": "СИМВ", "X": "ДР",
}

FEAT_RU = {
    "Animacy": "Одушевлённость", "Aspect": "Вид", "Case": "Падеж",
    "Degree": "Степень сравнения", "Gender": "Род", "Mood": "Наклонение",
    "Number": "Число", "Person": "Лицо", "Tense": "Время",
    "VerbForm": "Форма глагола", "Voice": "Залог",
    "Anim": "одушевлённое", "Inan": "неодушевлённое",
    "Imp": "несовершенный", "Perf": "совершенный",
    "Nom": "именительный", "Gen": "родительный", "Dat": "дательный",
    "Acc": "винительный", "Ins": "творительный", "Loc": "предложный",
    "Sing": "единственное", "Plur": "множественное",
    "Masc": "мужской", "Fem": "женский", "Neut": "средний",
    "Pres": "настоящее", "Past": "прошедшее", "Fut": "будущее",
    "Ind": "изъявительное", "Part": "причастие", "Conv": "деепричастие",
    "Inf": "инфинитив", "Fin": "финитная",
    "Act": "действительный", "Pass": "страдательный", "Mid": "средний",
}

STYLE_MARKERS = {
    "частицы": ["же", "ли", "вот", "ну", "уж", "даже", "лишь", "только"],
    "связки": ["в общем", "короче", "на самом деле", "так сказать", "как бы"],
    "союзы": ["потому что", "так как", "однако", "поэтому", "следовательно"],
}


@dataclass
class TokenInfo:
    text: str
    lemma: str
    pos: str         # UPOS tag
    pos_label: str   # русское название
    feats: str       # морфологические признаки


# ============================================================================
# STANZA-БЭКЕНД
# ============================================================================
class StanzaBackend:
    """Обёртка над Stanford Stanza для русского языка."""

    def __init__(self):
        self.nlp = None
        self._ready = False

    def ensure_loaded(self, status_callback=None):
        if self._ready:
            return
        import stanza
        # Патч для PyTorch 2.6+: weights_only=True по умолчанию ломает Stanza.
        # Безопасно переопределяем для моделей из официального репозитория Stanza.
        try:
            import torch
            _original_torch_load = torch.load
            def _patched_torch_load(*args, **kwargs):
                kwargs.setdefault("weights_only", False)
                return _original_torch_load(*args, **kwargs)
            torch.load = _patched_torch_load
        except Exception:
            pass
        if status_callback:
            status_callback("Загрузка модели Stanza (при первом запуске скачивается ~100 МБ)...")
        try:
            self.nlp = stanza.Pipeline('ru', processors='tokenize,pos,lemma', verbose=False)
        except Exception:
            if status_callback:
                status_callback("Скачивание модели русского языка...")
            stanza.download('ru', verbose=False)
            self.nlp = stanza.Pipeline('ru', processors='tokenize,pos,lemma', verbose=False)
        self._ready = True
        if status_callback:
            status_callback("Stanza готова")

    def analyze(self, text: str) -> List[TokenInfo]:
        if not self._ready:
            self.ensure_loaded()
        doc = self.nlp(text)
        tokens = []
        for sent in doc.sentences:
            for word in sent.words:
                if not WORD_RE.search(word.text):
                    tokens.append(TokenInfo(word.text, word.text, "PUNCT", "Пунктуация", "—"))
                    continue
                upos = word.upos or "X"
                feats_str = word.feats if word.feats else "—"
                feats_ru = self._translate_feats(feats_str) if feats_str != "—" else "—"
                pos_label = UPOS_RU.get(upos, upos)
                if upos == "VERB" and feats_str:
                    if "VerbForm=Part" in feats_str:
                        pos_label = "Причастие"
                    elif "VerbForm=Conv" in feats_str:
                        pos_label = "Деепричастие"
                tokens.append(TokenInfo(word.text, word.lemma or word.text.lower(),
                                       upos, pos_label, feats_ru))
        return tokens

    @staticmethod
    def _translate_feats(feats_str: str) -> str:
        if not feats_str or feats_str == "—":
            return "—"
        parts = []
        for pair in feats_str.split("|"):
            if "=" in pair:
                k, v = pair.split("=", 1)
                k_ru = FEAT_RU.get(k, k)
                v_ru = FEAT_RU.get(v, v)
                parts.append(f"{k_ru}: {v_ru}")
            else:
                parts.append(FEAT_RU.get(pair, pair))
        return "; ".join(parts)


# ============================================================================
# POS-БИГРАММЫ (Litvinova et al., 2015-2016)
# ============================================================================
def calculate_pos_bigrams(tokens: List[TokenInfo]) -> Dict:
    """
    Вычисление коэффициентов сочетаемости частеречных пар.

    Научная основа: Litvinova, Seredin & Litvinova (2015-2016) показали,
    что частотное распределение POS-биграмм в текстах позволяет оценивать
    индивидуальные характеристики авторов (пол, личностные черты).
    Исследовано 227 типов биграмм на русскоязычном корпусе "RusPersonality".

    Возвращает:
        - bigram_counts: Counter пар (POS1, POS2)
        - bigram_freq: нормированные частоты
        - matrix: матрица сочетаемости
        - top_bigrams: топ-20 биграмм
        - pos_labels: список POS-тегов (для осей матрицы)
    """
    words = [t for t in tokens if t.pos != "PUNCT" and WORD_RE.search(t.text)]
    if len(words) < 2:
        return {"bigram_counts": Counter(), "bigram_freq": {}, "matrix": {},
                "top_bigrams": [], "pos_labels": []}

    # Считаем биграммы POS-тегов
    bigrams = [(words[i].pos, words[i+1].pos) for i in range(len(words)-1)]
    bigram_counts = Counter(bigrams)
    total_bigrams = len(bigrams)

    # Нормированные частоты
    bigram_freq = {bg: round(cnt / total_bigrams, 5) for bg, cnt in bigram_counts.items()}

    # Матрица сочетаемости (все POS, встретившиеся в тексте)
    pos_set = sorted(set(t.pos for t in words))
    matrix = {}
    for p1 in pos_set:
        matrix[p1] = {}
        for p2 in pos_set:
            matrix[p1][p2] = bigram_freq.get((p1, p2), 0.0)

    # Топ-20 биграмм
    top = sorted(bigram_counts.items(), key=lambda x: -x[1])[:20]
    top_bigrams = [
        {
            "pair": bg,
            "pair_ru": f"{UPOS_SHORT.get(bg[0], bg[0])}→{UPOS_SHORT.get(bg[1], bg[1])}",
            "pair_full": f"{UPOS_RU.get(bg[0], bg[0])} → {UPOS_RU.get(bg[1], bg[1])}",
            "count": cnt,
            "freq": round(cnt / total_bigrams, 4),
        }
        for bg, cnt in top
    ]

    return {
        "bigram_counts": bigram_counts,
        "bigram_freq": bigram_freq,
        "matrix": matrix,
        "top_bigrams": top_bigrams,
        "pos_labels": pos_set,
        "total_bigrams": total_bigrams,
    }


# ============================================================================
# СТАТИСТИКА
# ============================================================================
def calculate_metrics(tokens: List[TokenInfo], text: str) -> dict:
    words = [t for t in tokens if WORD_RE.search(t.text) and t.pos != "PUNCT"]
    total = len(words)
    if total == 0:
        return {"частоты": {}, "дополнительно": {"Всего слов": 0},
                "профиль_служебных_слов": {}, "pos_bigrams": {}}

    pos_counts = Counter(t.pos_label for t in words)
    freq = {p: {"количество": c, "коэффициент": round(c / total, 4)}
            for p, c in sorted(pos_counts.items(), key=lambda x: (-x[1], x[0]))}

    lemmas = [t.lemma.lower() for t in words]
    sent_lens = [len(WORD_RE.findall(s)) for s in re.split(r"[.!?]+", text) if WORD_RE.findall(s)]
    word_lengths = [len(w.text) for w in words]

    content_set = {"Существительное", "Имя собственное", "Глагол", "Причастие",
                   "Деепричастие", "Прилагательное", "Наречие"}
    function_set = {"Предлог", "Частица", "Сочинительный союз", "Подчинительный союз",
                    "Местоимение", "Определительное слово"}
    content = sum(pos_counts.get(p, 0) for p in content_set)
    function = sum(pos_counts.get(p, 0) for p in function_set)
    noun_cnt = pos_counts.get("Существительное", 0) + pos_counts.get("Имя собственное", 0)
    verb_cnt = pos_counts.get("Глагол", 0) + pos_counts.get("Вспомогательный глагол", 0)

    additional = {
        "Всего слов": total,
        "Всего предложений": len(sent_lens),
        "Средняя длина слова (букв)": round(sum(word_lengths) / len(word_lengths), 2),
        "Лексическое разнообразие (TTR)": round(len(set(w.text.lower() for w in words)) / total, 4),
        "Лемматическое разнообразие": round(len(set(lemmas)) / total, 4),
        "Доля hapax-лемм": round(sum(1 for _, c in Counter(lemmas).items() if c == 1) / total, 4),
        "Средняя длина предложения (слов)": round(sum(sent_lens) / len(sent_lens), 2) if sent_lens else 0.0,
        "Дисперсия длины предложений": round(statistics.pvariance(sent_lens), 2) if len(sent_lens) > 1 else 0.0,
        "Содержательные/служебные": round(content / function, 4) if function else "∞",
        "Существительные/глаголы": round(noun_cnt / verb_cnt, 4) if verb_cnt else "∞",
    }

    lowered = text.lower()
    style_markers = {g: {m: len(re.findall(rf"(?<!\w){re.escape(m)}(?!\w)", lowered)) for m in marks}
                     for g, marks in STYLE_MARKERS.items()}

    # POS-биграммы
    pos_bigrams = calculate_pos_bigrams(tokens)

    return {
        "частоты": freq,
        "дополнительно": additional,
        "профиль_служебных_слов": style_markers,
        "pos_bigrams": pos_bigrams,
    }


# ============================================================================
# СРАВНЕНИЕ ТЕКСТОВ
# ============================================================================
def compare_texts(tokens1: List[TokenInfo], tokens2: List[TokenInfo],
                  text1: str, text2: str) -> dict:
    w1 = [t for t in tokens1 if t.pos != "PUNCT" and WORD_RE.search(t.text)]
    w2 = [t for t in tokens2 if t.pos != "PUNCT" and WORD_RE.search(t.text)]

    lemmas1 = [t.lemma.lower() for t in w1]
    lemmas2 = [t.lemma.lower() for t in w2]
    set1, set2 = set(lemmas1), set(lemmas2)
    common_lemmas = set1 & set2
    jaccard = len(common_lemmas) / len(set1 | set2) if (set1 | set2) else 0

    pos1 = Counter(t.pos_label for t in w1)
    pos2 = Counter(t.pos_label for t in w2)
    all_pos = set(pos1.keys()) | set(pos2.keys())
    tot1, tot2 = max(sum(pos1.values()), 1), max(sum(pos2.values()), 1)
    pos_diff = sum(abs(pos1.get(p, 0)/tot1 - pos2.get(p, 0)/tot2) for p in all_pos)
    pos_sim = max(0, 1 - pos_diff / 2)

    sl1 = [len(WORD_RE.findall(s)) for s in re.split(r'[.!?]+', text1) if WORD_RE.findall(s)]
    sl2 = [len(WORD_RE.findall(s)) for s in re.split(r'[.!?]+', text2) if WORD_RE.findall(s)]
    avg1 = sum(sl1)/len(sl1) if sl1 else 0
    avg2 = sum(sl2)/len(sl2) if sl2 else 0
    syn_sim = max(0, 1 - abs(avg1 - avg2) / max(avg1, avg2, 1))

    ttr1 = len(set1) / max(len(lemmas1), 1)
    ttr2 = len(set2) / max(len(lemmas2), 1)
    ttr_sim = max(0, 1 - abs(ttr1 - ttr2) / max(ttr1, ttr2, 0.01))

    # POS-биграммное сходство (новое!)
    bg1 = calculate_pos_bigrams(tokens1)
    bg2 = calculate_pos_bigrams(tokens2)
    bg_sim = _bigram_similarity(bg1["bigram_freq"], bg2["bigram_freq"])

    overall = jaccard * 0.30 + pos_sim * 0.20 + syn_sim * 0.15 + ttr_sim * 0.15 + bg_sim * 0.20

    service = {"и", "в", "не", "на", "с", "что", "а", "по", "это", "к", "но", "из",
               "у", "за", "о", "же", "от", "он", "она", "оно", "его", "её", "их",
               "все", "бы", "как", "мы", "вы", "они", "если", "или", "ни", "при",
               "до", "тот", "эта", "этот", "быть", "который", "свой", "весь"}
    meaningful = sorted(common_lemmas - service,
                        key=lambda x: -(Counter(lemmas1)[x] + Counter(lemmas2)[x]))[:30]

    return {
        "overall": round(overall, 3),
        "jaccard": round(jaccard, 3),
        "pos_similarity": round(pos_sim, 3),
        "syntactic_similarity": round(syn_sim, 3),
        "ttr_similarity": round(ttr_sim, 3),
        "bigram_similarity": round(bg_sim, 3),
        "common_lemmas": meaningful,
        "words1": len(w1), "words2": len(w2),
        "ttr1": round(ttr1, 4), "ttr2": round(ttr2, 4),
        "avg_sent1": round(avg1, 1), "avg_sent2": round(avg2, 1),
    }


def _bigram_similarity(freq1: dict, freq2: dict) -> float:
    """Косинусное сходство POS-биграммных профилей."""
    all_keys = set(freq1.keys()) | set(freq2.keys())
    if not all_keys:
        return 0.0
    dot = sum(freq1.get(k, 0) * freq2.get(k, 0) for k in all_keys)
    n1 = math.sqrt(sum(v*v for v in freq1.values())) or 1e-9
    n2 = math.sqrt(sum(v*v for v in freq2.values())) or 1e-9
    return dot / (n1 * n2)


# ============================================================================
# ЭКСПОРТ DOCX
# ============================================================================
def export_report_docx(filepath: str, text: str, metrics: dict,
                       error_result: ErrorAnalysisResult | None,
                       tokens: List[TokenInfo]):
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    h = doc.add_heading('СВОДНЫЙ ОТЧЁТ АВТОРОВЕДЧЕСКОГО АНАЛИЗА', level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'Дата: {datetime.now().strftime("%d.%m.%Y %H:%M")}')

    # 1. Описание
    doc.add_heading('1. Описание речевого продукта', level=2)
    words_list = re.findall(r"[А-Яа-яЁё]+", text)
    first_w = " ".join(words_list[:6]) + "..." if len(words_list) > 6 else " ".join(words_list)
    last_w = "..." + " ".join(words_list[-6:]) if len(words_list) > 6 else ""
    doc.add_paragraph(f'Текст на русском языке. Начало: «{first_w}»')
    if last_w:
        doc.add_paragraph(f'Окончание: «{last_w}»')
    doc.add_paragraph(f'Вербальный объём: {metrics["дополнительно"]["Всего слов"]} слов.')

    # 2. Количественные характеристики
    doc.add_heading('2. Количественные характеристики', level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    table.rows[0].cells[0].text = 'Показатель'
    table.rows[0].cells[1].text = 'Значение'
    for k, v in metrics["дополнительно"].items():
        row = table.add_row().cells
        row[0].text = str(k)
        row[1].text = str(v)

    # 3. Частоты частей речи
    doc.add_heading('3. Распределение частей речи', level=2)
    if metrics["частоты"]:
        t2 = doc.add_table(rows=1, cols=3)
        t2.style = 'Light Grid Accent 1'
        t2.rows[0].cells[0].text = 'Часть речи'
        t2.rows[0].cells[1].text = 'Количество'
        t2.rows[0].cells[2].text = 'Доля'
        for p, v in metrics["частоты"].items():
            row = t2.add_row().cells
            row[0].text = p
            row[1].text = str(v['количество'])
            row[2].text = f"{v['коэффициент']:.2%}"

    # 3.1. POS-биграммы
    pos_bg = metrics.get("pos_bigrams", {})
    top_bg = pos_bg.get("top_bigrams", [])
    if top_bg:
        doc.add_heading('3.1. Коэффициенты сочетаемости частеречных пар', level=2)
        doc.add_paragraph(
            'Метод POS-биграмм: частотное распределение последовательных пар '
            'частей речи отражает грамматические привычки автора и является '
            'идентификационным признаком (Litvinova et al., 2015–2016).'
        )
        t3 = doc.add_table(rows=1, cols=4)
        t3.style = 'Light Grid Accent 1'
        hdr = t3.rows[0].cells
        hdr[0].text = '№'
        hdr[1].text = 'Пара'
        hdr[2].text = 'Количество'
        hdr[3].text = 'Коэффициент'
        for i, bg in enumerate(top_bg[:15], 1):
            row = t3.add_row().cells
            row[0].text = str(i)
            row[1].text = bg["pair_full"]
            row[2].text = str(bg["count"])
            row[3].text = f'{bg["freq"]:.4f}'

    # 4. Навыки
    if error_result and error_result.skill_levels:
        doc.add_heading('4. Степени развития языковых навыков (по С.М. Вул)', level=2)
        for skill in error_result.skill_levels:
            p = doc.add_paragraph()
            run = p.add_run(f'{skill.skill_name}: ')
            run.bold = True
            p.add_run(f'{skill.level.upper()} ({skill.description})')

    # 5. Ошибки
    if error_result and error_result.errors:
        doc.add_heading('5. Выявленные речевые ошибки', level=2)
        by_type = {}
        for e in error_result.errors:
            by_type.setdefault(e.error_type, []).append(e)
        for etype, errs in by_type.items():
            doc.add_heading(f'{etype.upper()} ({len(errs)})', level=3)
            for i, e in enumerate(errs[:10], 1):
                txt = f'{i}. «{e.fragment}» — {e.description}'
                if e.suggestion:
                    txt += f' → {e.suggestion}'
                if e.rule_ref:
                    txt += f' [{e.rule_ref}]'
                doc.add_paragraph(txt)

    # 6. Вывод
    doc.add_heading('6. Вывод', level=2)
    wc = metrics["дополнительно"]["Всего слов"]
    if wc >= 500:
        doc.add_paragraph(f'Объём текста ({wc} слов) пригоден для судебной автороведческой экспертизы.')
    elif wc >= 200:
        doc.add_paragraph(f'Объём текста ({wc} слов) пригоден для предварительного анализа.')
    else:
        doc.add_paragraph(f'Объём текста ({wc} слов) недостаточен (минимум 500 слов).')
    doc.save(filepath)


def export_comparison_docx(filepath: str, comp: dict, text1: str, text2: str):
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    h = doc.add_heading('СРАВНИТЕЛЬНЫЙ АНАЛИЗ ТЕКСТОВ', level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'Дата: {datetime.now().strftime("%d.%m.%Y %H:%M")}')

    doc.add_heading('1. Общее сходство', level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    table.rows[0].cells[0].text = 'Компонент'
    table.rows[0].cells[1].text = 'Значение'
    for label, key in [("Общее сходство", "overall"), ("Лексическое (Jaccard)", "jaccard"),
                       ("Морфологическое (POS)", "pos_similarity"),
                       ("Синтаксическое", "syntactic_similarity"),
                       ("TTR-сходство", "ttr_similarity"),
                       ("POS-биграммное", "bigram_similarity")]:
        if key in comp:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = f"{comp[key]:.1%}"

    doc.add_heading('2. Совпадающие леммы', level=2)
    if comp["common_lemmas"]:
        doc.add_paragraph(", ".join(comp["common_lemmas"]))
    else:
        doc.add_paragraph("Совпадений не обнаружено.")

    doc.add_heading('3. Вывод', level=2)
    sim = comp["overall"]
    if sim >= 0.7:
        conclusion = "Высокая степень сходства. Возможна принадлежность одному автору."
    elif sim >= 0.5:
        conclusion = "Средняя степень сходства. Необходимо расширение материала."
    elif sim >= 0.3:
        conclusion = "Умеренные различия. Разные авторы или ситуации."
    else:
        conclusion = "Существенные различия. Вероятно, разные авторы."
    doc.add_paragraph(conclusion)
    doc.save(filepath)


# ============================================================================
# ЗАГРУЗКА ФАЙЛОВ
# ============================================================================
def load_text_from_file(filepath: str) -> str:
    ext = os.path.splitext(filepath)[1].lower()
    if ext == '.txt':
        for enc in ('utf-8', 'cp1251', 'cp866', 'latin-1'):
            try:
                with open(filepath, 'r', encoding=enc) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        raise ValueError(f"Не удалось прочитать: {filepath}")
    elif ext == '.docx':
        from docx import Document
        doc = Document(filepath)
        return '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
    else:
        raise ValueError(f"Неподдерживаемый формат: {ext}")


# ============================================================================
# КОНТЕКСТНОЕ МЕНЮ
# ============================================================================
class ContextMenuManager:
    """Менеджер контекстных меню для всех текстовых виджетов."""

    @staticmethod
    def attach(widget, readonly=False):
        """Привязать контекстное меню к текстовому виджету."""
        menu = tk.Menu(widget, tearoff=0)
        if not readonly:
            menu.add_command(label="✂ Вырезать  (Ctrl+X)",
                             command=lambda: ContextMenuManager._cut(widget))
            menu.add_command(label="📋 Вставить (Ctrl+V)",
                             command=lambda: ContextMenuManager._paste(widget))
        menu.add_command(label="📄 Копировать (Ctrl+C)",
                         command=lambda: ContextMenuManager._copy(widget))
        menu.add_separator()
        menu.add_command(label="🔘 Выделить всё (Ctrl+A)",
                         command=lambda: ContextMenuManager._select_all(widget))

        def _show_menu(event):
            menu.tk_popup(event.x_root, event.y_root)

        widget.bind("<Button-3>", _show_menu)

        # Горячие клавиши (для editable виджетов они уже есть в Tk, но для
        # ScrolledText в состоянии disabled — нужно явно)
        if not readonly:
            widget.bind("<Control-x>", lambda e: ContextMenuManager._cut(widget))
            widget.bind("<Control-v>", lambda e: ContextMenuManager._paste(widget))
        widget.bind("<Control-c>", lambda e: ContextMenuManager._copy(widget))
        widget.bind("<Control-a>", lambda e: ContextMenuManager._select_all(widget))
        # Примечание: Tkinter на Windows не поддерживает кириллические keysym,
        # поэтому Ctrl+C/V/X в русской раскладке обрабатываются ОС автоматически.

    @staticmethod
    def _cut(widget):
        try:
            if widget.tag_ranges("sel"):
                widget.event_generate("<<Cut>>")
        except tk.TclError:
            pass

    @staticmethod
    def _copy(widget):
        try:
            if widget.tag_ranges("sel"):
                widget.event_generate("<<Copy>>")
        except tk.TclError:
            pass

    @staticmethod
    def _paste(widget):
        try:
            widget.event_generate("<<Paste>>")
        except tk.TclError:
            pass

    @staticmethod
    def _select_all(widget):
        widget.tag_add("sel", "1.0", "end")
        return "break"


# ============================================================================
# ТЕПЛОВАЯ КАРТА POS-БИГРАММ
# ============================================================================
class HeatmapWindow:
    """Окно с тепловой картой сочетаемости частеречных пар."""

    def __init__(self, parent, pos_bigrams: dict, title="Тепловая карта POS-биграмм"):
        self.win = tk.Toplevel(parent)
        self.win.title(title)
        self.win.geometry("800x700")

        matrix = pos_bigrams.get("matrix", {})
        labels = pos_bigrams.get("pos_labels", [])
        if not labels or not matrix:
            ttk.Label(self.win, text="Нет данных для отображения").pack(pady=20)
            return

        # Заголовок
        ttk.Label(self.win, text="Коэффициенты сочетаемости частеречных пар (POS-биграммы)",
                  font=("Arial", 11, "bold")).pack(pady=(8, 2))
        ttk.Label(self.win, text="Строки — первый элемент пары, столбцы — второй",
                  font=("Arial", 9)).pack(pady=(0, 4))

        # Canvas для рисования
        canvas_frame = ttk.Frame(self.win)
        canvas_frame.pack(fill="both", expand=True, padx=8, pady=4)

        canvas = tk.Canvas(canvas_frame, bg="white")
        canvas.pack(fill="both", expand=True)

        self._draw_heatmap(canvas, matrix, labels)

    def _draw_heatmap(self, canvas, matrix, labels):
        """Рисование тепловой карты на Canvas."""
        n = len(labels)
        if n == 0:
            return

        short_labels = [UPOS_SHORT.get(l, l[:4]) for l in labels]

        margin_left = 70
        margin_top = 70
        margin_right = 120  # для легенды
        margin_bottom = 20

        canvas.update_idletasks()
        w = max(canvas.winfo_width(), 600)
        h = max(canvas.winfo_height(), 500)

        cell_w = min(45, (w - margin_left - margin_right) // max(n, 1))
        cell_h = min(35, (h - margin_top - margin_bottom) // max(n, 1))

        # Найти максимум для нормировки цвета
        max_val = 0
        for p1 in labels:
            for p2 in labels:
                val = matrix.get(p1, {}).get(p2, 0)
                if val > max_val:
                    max_val = val
        max_val = max(max_val, 0.001)

        # Метки столбцов (сверху)
        for j, label in enumerate(short_labels):
            x = margin_left + j * cell_w + cell_w // 2
            canvas.create_text(x, margin_top - 8, text=label, angle=45,
                               anchor="s", font=("Arial", 7, "bold"))

        # Рисуем ячейки
        for i, p1 in enumerate(labels):
            # Метки строк (слева)
            y = margin_top + i * cell_h + cell_h // 2
            canvas.create_text(margin_left - 5, y, text=short_labels[i],
                               anchor="e", font=("Arial", 7, "bold"))

            for j, p2 in enumerate(labels):
                val = matrix.get(p1, {}).get(p2, 0)
                x1 = margin_left + j * cell_w
                y1 = margin_top + i * cell_h
                x2 = x1 + cell_w
                y2 = y1 + cell_h

                # Цвет: от белого (0) к тёмно-синему (max)
                intensity = val / max_val if max_val > 0 else 0
                r = int(255 * (1 - intensity * 0.85))
                g = int(255 * (1 - intensity * 0.7))
                b = 255
                color = f"#{r:02x}{g:02x}{b:02x}"

                canvas.create_rectangle(x1, y1, x2, y2, fill=color, outline="#ccc")

                # Число в ячейке (если достаточно места)
                if cell_w >= 25 and val > 0:
                    txt_color = "white" if intensity > 0.5 else "#333"
                    canvas.create_text((x1+x2)//2, (y1+y2)//2,
                                       text=f"{val:.3f}" if val < 0.1 else f"{val:.2f}",
                                       font=("Arial", 6), fill=txt_color)

        # Легенда
        leg_x = margin_left + n * cell_w + 20
        leg_y = margin_top
        canvas.create_text(leg_x, leg_y - 10, text="Легенда", anchor="w",
                           font=("Arial", 9, "bold"))
        leg_height = min(200, n * cell_h)
        for i in range(20):
            frac = i / 19
            r = int(255 * (1 - frac * 0.85))
            g = int(255 * (1 - frac * 0.7))
            b = 255
            c = f"#{r:02x}{g:02x}{b:02x}"
            yy = leg_y + int(frac * leg_height)
            canvas.create_rectangle(leg_x, yy, leg_x + 20, yy + leg_height // 20 + 1,
                                    fill=c, outline="")
        canvas.create_text(leg_x + 25, leg_y, text="0", anchor="w", font=("Arial", 7))
        canvas.create_text(leg_x + 25, leg_y + leg_height,
                           text=f"{max_val:.3f}", anchor="w", font=("Arial", 7))


# ============================================================================
# ГРАФИЧЕСКИЙ ИНТЕРФЕЙС
# ============================================================================
class AutorovedApp:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("Автороведческий анализатор v4 (Stanza + POS-биграммы)")
        self.root.geometry("1400x950")

        self.stanza = StanzaBackend()
        self.error_analyzer = ErrorAnalyzer()
        self._last_tokens: List[TokenInfo] = []
        self._last_freq: dict = {}
        self._last_text: str = ""
        self._last_metrics: dict = {}
        self._last_error_result: ErrorAnalysisResult | None = None
        self._last_strat_result = None
        self._row_to_span: dict = {}
        self._stanza_loaded = False

        self._build_ui()
        self._bind_hotkeys()
        self._attach_context_menus()

    # ───── UI СТРОИТЕЛЬ ─────
    def _build_ui(self):
        top = ttk.Frame(self.root, padding=8)
        top.pack(fill="x")
        ttk.Label(top, text="Введите текст для анализа или загрузите файл (.txt / .docx):",
                  font=("Arial", 10, "bold")).pack(anchor="w")

        ctrl = ttk.Frame(top)
        ctrl.pack(fill="x", pady=2)
        ttk.Button(ctrl, text="📂 Загрузить файл", command=self._load_file).pack(side="left")
        self.status_var = tk.StringVar(value="Stanza: не загружена (загрузится при первом анализе)")
        ttk.Label(ctrl, textvariable=self.status_var, foreground="#666").pack(side="right")

        self.text_input = tk.Text(top, height=10, wrap="word", font=("Consolas", 10))
        self.text_input.pack(fill="x", pady=(2, 6))
        self.text_input.tag_configure("hover_token", background="#fff176")
        self.text_input.tag_configure("error_highlight", background="#ffcdd2")

        btns = ttk.Frame(top)
        btns.pack(fill="x")
        ttk.Button(btns, text="▶ Полный анализ (Ctrl+Enter)", command=self._run_analysis_thread).pack(side="left")
        ttk.Button(btns, text="📈 Диаграмма POS", command=self.show_pie_chart).pack(side="left", padx=5)
        ttk.Button(btns, text="🔥 Heatmap биграмм", command=self.show_heatmap).pack(side="left", padx=5)
        ttk.Button(btns, text="🗑 Очистить", command=self.clear_all).pack(side="left", padx=5)

        # Notebook
        self.notebook = ttk.Notebook(self.root)
        self.notebook.pack(fill="both", expand=True, padx=8, pady=4)

        # Вкладка 1: Морфология
        self.tab_morph = ttk.Frame(self.notebook)
        self.notebook.add(self.tab_morph, text="Морфология")
        self._build_morph_tab()

        # Вкладка 2: Статистика
        self.tab_stats = ttk.Frame(self.notebook)
        self.notebook.add(self.tab_stats, text="Статистика")
        self.stats_text = scrolledtext.ScrolledText(self.tab_stats, wrap="word", font=("Consolas", 10))
        self.stats_text.pack(fill="both", expand=True, padx=4, pady=4)
        self.stats_text.configure(state="disabled")

        # Вкладка 3: Ошибки и навыки
        self.tab_errors = ttk.Frame(self.notebook)
        self.notebook.add(self.tab_errors, text="Ошибки и навыки")
        self._build_errors_tab()

        # Вкладка 4: Интернет-коммуникация
        self.tab_internet = ttk.Frame(self.notebook)
        self.notebook.add(self.tab_internet, text="Интернет-коммуникация")
        self.internet_score_label = ttk.Label(self.tab_internet, text="—", font=("Arial", 12, "bold"))
        self.internet_score_label.pack(padx=10, pady=6)
        self.internet_text = scrolledtext.ScrolledText(self.tab_internet, wrap="word", font=("Consolas", 10))
        self.internet_text.pack(fill="both", expand=True, padx=4, pady=4)
        self.internet_text.configure(state="disabled")

        # Вкладка 5: Лексическая стратификация
        self.tab_strat = ttk.Frame(self.notebook)
        self.notebook.add(self.tab_strat, text="Стратификация лексики")
        self._build_strat_tab()

        # Вкладка 6: Сравнение текстов
        self.tab_compare = ttk.Frame(self.notebook)
        self.notebook.add(self.tab_compare, text="Сравнение текстов")
        self._build_compare_tab()

        # Вкладка 7: Сводный отчёт
        self.tab_report = ttk.Frame(self.notebook)
        self.notebook.add(self.tab_report, text="Сводный отчёт")
        self.report_text = scrolledtext.ScrolledText(self.tab_report, wrap="word", font=("Consolas", 10))
        self.report_text.pack(fill="both", expand=True, padx=4, pady=4)
        self.report_text.configure(state="disabled")
        btn_exp = ttk.Frame(self.tab_report)
        btn_exp.pack(fill="x", padx=4, pady=4)
        ttk.Button(btn_exp, text="📄 Экспорт отчёта в DOCX", command=self._export_docx).pack(side="left")

    def _build_morph_tab(self):
        self.tokens_table = ttk.Treeview(self.tab_morph,
            columns=("token", "lemma", "pos", "feats"), show="headings", height=20)
        for col, w, title in [("token", 150, "Словоформа"), ("lemma", 150, "Лемма"),
                               ("pos", 150, "Часть речи"), ("feats", 500, "Морф. признаки")]:
            self.tokens_table.heading(col, text=title)
            self.tokens_table.column(col, width=w, anchor="w")
        scr = ttk.Scrollbar(self.tab_morph, orient="vertical", command=self.tokens_table.yview)
        self.tokens_table.configure(yscrollcommand=scr.set)
        self.tokens_table.pack(side="left", fill="both", expand=True)
        scr.pack(side="right", fill="y")
        self.tokens_table.bind("<Motion>", self._on_table_hover)

    def _build_errors_tab(self):
        sf = ttk.LabelFrame(self.tab_errors, text="Степени развития навыков (по С.М. Вул)")
        sf.pack(fill="x", padx=4, pady=4)
        self.skills_labels = {}
        grid = ttk.Frame(sf)
        grid.pack(fill="x", padx=8, pady=6)
        for i, name in enumerate(["Пунктуационные", "Орфографические", "Грамматические", "Лексико-фразеологические"]):
            ttk.Label(grid, text=f"{name} навыки:", font=("Arial", 9, "bold")).grid(row=i, column=0, sticky="w", pady=2)
            self.skills_labels[name] = ttk.Label(grid, text="—", font=("Arial", 9))
            self.skills_labels[name].grid(row=i, column=1, sticky="w", padx=10)

        ef = ttk.LabelFrame(self.tab_errors, text="Обнаруженные ошибки")
        ef.pack(fill="both", expand=True, padx=4, pady=4)
        self.errors_tree = ttk.Treeview(ef,
            columns=("type", "subtype", "fragment", "description", "suggestion"), show="headings")
        for col, w, title in [("type", 120, "Тип"), ("subtype", 130, "Подтип"),
                               ("fragment", 150, "Фрагмент"), ("description", 300, "Описание"),
                               ("suggestion", 200, "Рекомендация")]:
            self.errors_tree.heading(col, text=title)
            self.errors_tree.column(col, width=w, anchor="w")
        escr = ttk.Scrollbar(ef, orient="vertical", command=self.errors_tree.yview)
        self.errors_tree.configure(yscrollcommand=escr.set)
        self.errors_tree.pack(side="left", fill="both", expand=True)
        escr.pack(side="right", fill="y")
        self.errors_tree.bind("<<TreeviewSelect>>", self._on_error_select)

    def _build_strat_tab(self):
        """Вкладка лексической стратификации."""
        top_frame = ttk.LabelFrame(self.tab_strat, text="Распределение по стилистическим пластам")
        top_frame.pack(fill="x", padx=6, pady=4)

        cols = ("layer", "count", "ratio", "examples")
        self.strat_tree = ttk.Treeview(top_frame, columns=cols, show="headings", height=9)
        for col, w, title in [
            ("layer", 220, "Стилистический пласт"),
            ("count", 60,  "Ед."),
            ("ratio", 70,  "Доля"),
            ("examples", 450, "Примеры"),
        ]:
            self.strat_tree.heading(col, text=title)
            self.strat_tree.column(col, width=w, anchor="w")
        strat_scr = ttk.Scrollbar(top_frame, orient="vertical", command=self.strat_tree.yview)
        self.strat_tree.configure(yscrollcommand=strat_scr.set)
        self.strat_tree.pack(side="left", fill="both", expand=True)
        strat_scr.pack(side="right", fill="y")

        bottom_frame = ttk.LabelFrame(self.tab_strat, text="Подробный отчёт и аналитический вывод")
        bottom_frame.pack(fill="both", expand=True, padx=6, pady=4)
        self.strat_text = scrolledtext.ScrolledText(
            bottom_frame, wrap="word", font=("Consolas", 10), height=12)
        self.strat_text.pack(fill="both", expand=True, padx=4, pady=4)
        self.strat_text.configure(state="disabled")

        status_text = ("✓ Словарь загружен (12 251 лемма)"
                       if _STRATIFICATION_AVAILABLE
                       else "⚠ Файлы lexical_stratification.py и lexicon_stratified.json не найдены")
        ttk.Label(self.tab_strat, text=status_text, font=("Arial", 8),
                  foreground="#4caf50" if _STRATIFICATION_AVAILABLE else "#e53935"
                  ).pack(anchor="w", padx=6, pady=2)

    def _build_compare_tab(self):
        pane = ttk.PanedWindow(self.tab_compare, orient="horizontal")
        pane.pack(fill="both", expand=True, padx=4, pady=4)

        left = ttk.Frame(pane)
        pane.add(left, weight=1)

        ttk.Label(left, text="Текст 1:", font=("Arial", 9, "bold")).pack(anchor="w")
        f1 = ttk.Frame(left)
        f1.pack(fill="x")
        ttk.Button(f1, text="📂", command=lambda: self._load_compare(1)).pack(side="left")
        self.cmp_text1 = tk.Text(left, height=8, wrap="word", font=("Consolas", 9))
        self.cmp_text1.pack(fill="both", expand=True, pady=2)

        ttk.Label(left, text="Текст 2:", font=("Arial", 9, "bold")).pack(anchor="w")
        f2 = ttk.Frame(left)
        f2.pack(fill="x")
        ttk.Button(f2, text="📂", command=lambda: self._load_compare(2)).pack(side="left")
        self.cmp_text2 = tk.Text(left, height=8, wrap="word", font=("Consolas", 9))
        self.cmp_text2.pack(fill="both", expand=True, pady=2)

        ttk.Button(left, text="▶ Сравнить тексты", command=self._run_compare_thread).pack(pady=4)

        right = ttk.Frame(pane)
        pane.add(right, weight=1)
        self.cmp_result_text = scrolledtext.ScrolledText(right, wrap="word", font=("Consolas", 10))
        self.cmp_result_text.pack(fill="both", expand=True)
        self.cmp_result_text.configure(state="disabled")
        ttk.Button(right, text="📄 Экспорт сравнения в DOCX",
                   command=self._export_compare_docx).pack(pady=4)
        self._last_comparison = None

    # ───── КОНТЕКСТНЫЕ МЕНЮ ─────
    def _attach_context_menus(self):
        """Привязать контекстные меню ко всем текстовым виджетам."""
        # Редактируемые поля
        ContextMenuManager.attach(self.text_input, readonly=False)
        ContextMenuManager.attach(self.cmp_text1, readonly=False)
        ContextMenuManager.attach(self.cmp_text2, readonly=False)
        # Только чтение
        ContextMenuManager.attach(self.stats_text, readonly=True)
        ContextMenuManager.attach(self.internet_text, readonly=True)
        ContextMenuManager.attach(self.report_text, readonly=True)
        ContextMenuManager.attach(self.cmp_result_text, readonly=True)

    # ───── ЗАГРУЗКА ФАЙЛА ─────
    def _load_file(self):
        fp = filedialog.askopenfilename(filetypes=[("Тексты", "*.txt *.docx"), ("Все файлы", "*.*")])
        if fp:
            try:
                text = load_text_from_file(fp)
                self.text_input.delete("1.0", "end")
                self.text_input.insert("1.0", text)
            except Exception as e:
                messagebox.showerror("Ошибка", str(e))

    def _load_compare(self, num):
        fp = filedialog.askopenfilename(filetypes=[("Тексты", "*.txt *.docx")])
        if fp:
            try:
                text = load_text_from_file(fp)
                widget = self.cmp_text1 if num == 1 else self.cmp_text2
                widget.delete("1.0", "end")
                widget.insert("1.0", text)
            except Exception as e:
                messagebox.showerror("Ошибка", str(e))

    # ───── АНАЛИЗ ─────
    def _run_analysis_thread(self):
        text = self.text_input.get("1.0", "end").strip()
        if not text:
            messagebox.showwarning("Нет текста", "Введите текст для анализа.")
            return
        self.status_var.set("Анализ...")
        self.root.update_idletasks()
        threading.Thread(target=self._do_analysis, args=(text,), daemon=True).start()

    def _do_analysis(self, text):
        try:
            self.stanza.ensure_loaded(lambda msg: self.root.after(0, self.status_var.set, msg))
            tokens = self.stanza.analyze(text)
            self.root.after(0, self._finish_analysis, text, tokens)
        except Exception as e:
            self.root.after(0, lambda: messagebox.showerror("Ошибка", str(e)))
            self.root.after(0, self.status_var.set, f"Ошибка: {e}")

    def _finish_analysis(self, text, tokens):
        self._last_text = text
        self._last_tokens = tokens
        self._stanza_loaded = True
        self.status_var.set(f"Stanza: готово | {len([t for t in tokens if t.pos != 'PUNCT'])} слов")

        # Морфология
        for r in self.tokens_table.get_children():
            self.tokens_table.delete(r)
        self._map_rows(tokens, text)

        # Статистика + POS-биграммы
        metrics = calculate_metrics(tokens, text)
        self._last_metrics = metrics
        self._last_freq = metrics["частоты"]
        self._update_stats(metrics)

        # Ошибки
        result = self.error_analyzer.analyze(text)
        self._last_error_result = result
        self._update_errors(result)
        self._update_internet(result.internet_profile)

        # Лексическая стратификация
        if _STRATIFICATION_AVAILABLE:
            lemma_pairs = [(t.text, t.lemma) for t in tokens
                           if WORD_RE.search(t.text) and t.pos != "PUNCT"]
            strat_result = analyze_stratification(text, lemma_pairs)
            self._last_strat_result = strat_result
            self._update_strat(strat_result)
        else:
            self._last_strat_result = None

        # Отчёт
        self._generate_report(text, metrics, result)
        self.notebook.select(self.tab_stats)

    def _map_rows(self, tokens, text):
        self._row_to_span.clear()
        lower = text.lower()
        cursor = 0
        for tok in tokens:
            if not WORD_RE.search(tok.text):
                continue
            found = lower.find(tok.text.lower(), cursor)
            if found == -1:
                found = lower.find(tok.text.lower())
                if found == -1:
                    continue
            row = self.tokens_table.insert("", "end",
                values=(tok.text, tok.lemma, tok.pos_label, tok.feats))
            self._row_to_span[row] = (found, found + len(tok.text))
            cursor = found + len(tok.text)

    def _update_stats(self, metrics):
        lines = ["=" * 60, "ЛИНГВОСТАТИСТИЧЕСКИЕ ПОКАЗАТЕЛИ", "=" * 60, ""]
        for k, v in metrics["дополнительно"].items():
            lines.append(f"  {k}: {v}")

        lines.append("\n--- Частоты частей речи ---")
        for p, v in metrics["частоты"].items():
            lines.append(f"  {p}: {v['количество']} ({v['коэффициент']:.2%})")

        lines.append("\n--- Профиль служебных слов ---")
        for g, vals in metrics["профиль_служебных_слов"].items():
            found = [f"{k}={v}" for k, v in vals.items() if v > 0]
            if found:
                lines.append(f"  {g}: " + ", ".join(found))

        # POS-биграммы
        pos_bg = metrics.get("pos_bigrams", {})
        top_bg = pos_bg.get("top_bigrams", [])
        if top_bg:
            lines.append("\n" + "=" * 60)
            lines.append("КОЭФФИЦИЕНТЫ СОЧЕТАЕМОСТИ ЧАСТЕРЕЧНЫХ ПАР (POS-биграммы)")
            lines.append("=" * 60)
            lines.append(f"  Всего биграмм: {pos_bg.get('total_bigrams', 0)}")
            lines.append(f"  Уникальных типов: {len(pos_bg.get('bigram_freq', {}))}")
            lines.append("")
            lines.append(f"  {'№':>3}  {'Пара':<35} {'Кол-во':>8}  {'Коэфф.':>10}")
            lines.append("  " + "-" * 60)
            for i, bg in enumerate(top_bg, 1):
                lines.append(f"  {i:>3}. {bg['pair_full']:<35} {bg['count']:>8}  {bg['freq']:>10.4f}")

            lines.append("")
            lines.append("  Методическое обоснование:")
            lines.append("  POS-биграммы отражают грамматические привычки автора")
            lines.append("  при построении фраз. Litvinova et al. (2015-2016)")
            lines.append("  установили корреляции POS-биграмм с полом и личностными")
            lines.append("  чертами авторов на русскоязычном корпусе.")
            lines.append("  Кнопка «🔥 Heatmap биграмм» — тепловая карта сочетаемости.")

        self._set_text(self.stats_text, "\n".join(lines))

    def _update_errors(self, result):
        for r in self.errors_tree.get_children():
            self.errors_tree.delete(r)
        skill_map = {"Пунктуационные навыки": "Пунктуационные",
                     "Орфографические навыки": "Орфографические",
                     "Грамматические навыки": "Грамматические",
                     "Лексико-фразеологические навыки": "Лексико-фразеологические"}
        colors = {"высокая": "#4caf50", "средняя": "#ff9800", "низкая": "#f44336", "нулевая": "#9c27b0"}
        for skill in result.skill_levels:
            key = skill_map.get(skill.skill_name, skill.skill_name)
            if key in self.skills_labels:
                self.skills_labels[key].configure(
                    text=f"{skill.level.upper()} ({skill.description})",
                    foreground=colors.get(skill.level, "#000"))
        for error in result.errors:
            self.errors_tree.insert("", "end", values=(
                error.error_type, error.subtype,
                error.fragment[:30] + "..." if len(error.fragment) > 30 else error.fragment,
                error.description[:50] + "..." if len(error.description) > 50 else error.description,
                error.suggestion[:30] + "..." if len(error.suggestion) > 30 else error.suggestion))

    def _update_internet(self, profile):
        score = round(profile.internet_comm_score * 100, 1)
        assess = ("ВЫСОКАЯ" if score >= 50 else "СРЕДНЯЯ" if score >= 25
                  else "НИЗКАЯ" if score >= 10 else "МИНИМАЛЬНАЯ")
        col = "#e91e63" if score >= 50 else "#ff9800" if score >= 25 else "#2196f3" if score >= 10 else "#4caf50"
        self.internet_score_label.configure(text=f"{score}% — {assess} степень", foreground=col)
        lines = [format_internet_profile(profile)]
        self._set_text(self.internet_text, "\n".join(lines))

    def _update_strat(self, result: "StratificationResult"):
        """Обновляем вкладку лексической стратификации."""
        # Очищаем таблицу
        for row in self.strat_tree.get_children():
            self.strat_tree.delete(row)

        # Заполняем таблицу по слоям
        for layer in LAYER_ORDER:
            count = result.layer_counts.get(layer, 0)
            if count == 0:
                continue
            ratio = result.layer_ratios.get(layer, 0.0)
            label = LAYER_LABELS.get(layer, layer)
            tokens = result.by_layer.get(layer, [])
            # топ-5 уникальных примеров
            seen: dict = {}
            for t in tokens:
                key = t.match_key if t.is_phrase else t.lemma
                seen[key] = seen.get(key, 0) + 1
            top = sorted(seen, key=lambda x: -seen[x])[:5]
            examples = ", ".join(top)
            self.strat_tree.insert("", "end", values=(
                label,
                count,
                f"{ratio:.1%}",
                examples,
            ))

        # Подробный текстовый отчёт
        if _STRATIFICATION_AVAILABLE:
            report = format_stratification_report(result)
        else:
            report = "Модуль стратификации недоступен."
        self._set_text(self.strat_text, report)

    def _generate_report(self, text, metrics, error_result):
        lines = ["=" * 70, "СВОДНЫЙ ОТЧЁТ АВТОРОВЕДЧЕСКОГО АНАЛИЗА", "=" * 70, ""]
        words_l = re.findall(r"[А-Яа-яЁё]+", text)
        first_w = " ".join(words_l[:5]) + "..." if len(words_l) > 5 else " ".join(words_l)
        last_w = "..." + " ".join(words_l[-5:]) if len(words_l) > 5 else ""
        lines.append("1. ОПИСАНИЕ РЕЧЕВОГО ПРОДУКТА")
        lines.append(f"   Начало: «{first_w}»")
        if last_w:
            lines.append(f"   Конец: «{last_w}»")
        lines.append(f"   Объём: {metrics['дополнительно']['Всего слов']} слов")
        lines.append("\n2. КОЛИЧЕСТВЕННЫЕ ХАРАКТЕРИСТИКИ")
        for k, v in metrics["дополнительно"].items():
            lines.append(f"   {k}: {v}")

        # POS-биграммы в отчёте
        pos_bg = metrics.get("pos_bigrams", {})
        top_bg = pos_bg.get("top_bigrams", [])
        if top_bg:
            lines.append("\n2.1. ЧАСТЕРЕЧНЫЕ БИГРАММЫ (топ-10)")
            for i, bg in enumerate(top_bg[:10], 1):
                lines.append(f"   {i}. {bg['pair_full']}: {bg['count']} ({bg['freq']:.4f})")

        if error_result:
            lines.append("\n3. СТЕПЕНИ РАЗВИТИЯ НАВЫКОВ (по С.М. Вул)")
            for s in error_result.skill_levels:
                lines.append(f"   {s.skill_name}: {s.level.upper()} ({s.description})")
            lines.append(f"\n4. ОШИБКИ: {len(error_result.errors)}")
            by_type = {}
            for e in error_result.errors:
                by_type.setdefault(e.error_type, []).append(e)
            for et, errs in by_type.items():
                lines.append(f"   {et}: {len(errs)}")
                for i, e in enumerate(errs[:5], 1):
                    ref = f" [{e.rule_ref}]" if e.rule_ref else ""
                    lines.append(f"     {i}. «{e.fragment}» — {e.description}{ref}")

        # Лексическая стратификация в сводном отчёте
        if _STRATIFICATION_AVAILABLE and self._last_strat_result:
            sr = self._last_strat_result
            lines.append("\n5. ЛЕКСИЧЕСКАЯ СТРАТИФИКАЦИЯ (ЭКЦ МВД России)")
            lines.append(f"   Маркированных единиц: {sr.marked_words} из {sr.total_words} слов")
            for layer in LAYER_ORDER:
                cnt = sr.layer_counts.get(layer, 0)
                if cnt == 0:
                    continue
                ratio = sr.layer_ratios.get(layer, 0.0)
                label = LAYER_LABELS.get(layer, layer)
                lines.append(f"   {label}: {cnt} ({ratio:.1%})")
            from lexical_stratification import _make_conclusion
            lines.append("\n   Вывод по стратификации:")
            for l in _make_conclusion(sr).split("\n"):
                lines.append(f"  {l}")

        wc = metrics["дополнительно"]["Всего слов"]
        lines.append("\n6. ВЫВОД")
        if wc >= 500:
            lines.append("   Текст ПРИГОДЕН для судебной автороведческой экспертизы.")
        elif wc >= 200:
            lines.append("   Текст пригоден для предварительного анализа.")
        else:
            lines.append("   Текст НЕДОСТАТОЧЕН для экспертизы (мин. 500 слов).")
        self._set_text(self.report_text, "\n".join(lines))

    # ───── СРАВНЕНИЕ ─────
    def _run_compare_thread(self):
        t1 = self.cmp_text1.get("1.0", "end").strip()
        t2 = self.cmp_text2.get("1.0", "end").strip()
        if not t1 or not t2:
            messagebox.showwarning("Нет текстов", "Введите оба текста для сравнения.")
            return
        self.status_var.set("Сравнение текстов...")
        threading.Thread(target=self._do_compare, args=(t1, t2), daemon=True).start()

    def _do_compare(self, t1, t2):
        try:
            self.stanza.ensure_loaded(lambda msg: self.root.after(0, self.status_var.set, msg))
            tok1 = self.stanza.analyze(t1)
            tok2 = self.stanza.analyze(t2)
            comp = compare_texts(tok1, tok2, t1, t2)
            self.root.after(0, self._finish_compare, comp, t1, t2)
        except Exception as e:
            self.root.after(0, lambda: messagebox.showerror("Ошибка", str(e)))

    def _finish_compare(self, comp, t1, t2):
        self._last_comparison = comp
        self._last_cmp_texts = (t1, t2)
        self.status_var.set("Сравнение завершено")

        sim = comp["overall"]
        if sim >= 0.7:
            conclusion = "Высокое сходство — возможно один автор."
        elif sim >= 0.5:
            conclusion = "Среднее сходство — необходим дополнительный материал."
        elif sim >= 0.3:
            conclusion = "Умеренные различия — разные авторы или ситуации."
        else:
            conclusion = "Существенные различия — вероятно, разные авторы."

        lines = ["=" * 60, "РЕЗУЛЬТАТ СРАВНЕНИЯ ТЕКСТОВ", "=" * 60, ""]
        lines.append(f"  ОБЩЕЕ СХОДСТВО: {comp['overall']:.1%}")
        lines.append(f"  Лексическое (Jaccard): {comp['jaccard']:.1%}")
        lines.append(f"  Морфологическое (POS): {comp['pos_similarity']:.1%}")
        lines.append(f"  Синтаксическое: {comp['syntactic_similarity']:.1%}")
        lines.append(f"  TTR-сходство: {comp['ttr_similarity']:.1%}")
        lines.append(f"  POS-биграммное: {comp.get('bigram_similarity', 0):.1%}")
        lines.append("")
        lines.append(f"  Текст 1: {comp['words1']} слов, TTR={comp['ttr1']}, ср.предл.={comp['avg_sent1']}")
        lines.append(f"  Текст 2: {comp['words2']} слов, TTR={comp['ttr2']}, ср.предл.={comp['avg_sent2']}")
        lines.append("")
        lines.append(f"  Совпадающие леммы ({len(comp['common_lemmas'])}):")
        if comp["common_lemmas"]:
            for i in range(0, len(comp["common_lemmas"]), 6):
                chunk = comp["common_lemmas"][i:i+6]
                lines.append("    " + ", ".join(chunk))
        lines.append("")
        lines.append(f"  ВЫВОД: {conclusion}")
        self._set_text(self.cmp_result_text, "\n".join(lines))

    # ───── ЭКСПОРТ ─────
    def _export_docx(self):
        if not self._last_metrics:
            messagebox.showinfo("Нет данных", "Сначала выполните анализ.")
            return
        fp = filedialog.asksaveasfilename(defaultextension=".docx",
                                          filetypes=[("Word", "*.docx")],
                                          initialfile="отчёт_автороведческий.docx")
        if fp:
            try:
                export_report_docx(fp, self._last_text, self._last_metrics,
                                   self._last_error_result, self._last_tokens)
                messagebox.showinfo("Готово", f"Отчёт сохранён:\n{fp}")
            except Exception as e:
                messagebox.showerror("Ошибка", str(e))

    def _export_compare_docx(self):
        if not self._last_comparison:
            messagebox.showinfo("Нет данных", "Сначала выполните сравнение.")
            return
        fp = filedialog.asksaveasfilename(defaultextension=".docx",
                                          filetypes=[("Word", "*.docx")],
                                          initialfile="сравнение_текстов.docx")
        if fp:
            try:
                export_comparison_docx(fp, self._last_comparison, *self._last_cmp_texts)
                messagebox.showinfo("Готово", f"Отчёт сохранён:\n{fp}")
            except Exception as e:
                messagebox.showerror("Ошибка", str(e))

    # ───── ВСПОМОГАТЕЛЬНОЕ ─────
    def _on_table_hover(self, event):
        row_id = self.tokens_table.identify_row(event.y)
        self.text_input.tag_remove("hover_token", "1.0", "end")
        if row_id in self._row_to_span:
            a, b = self._row_to_span[row_id]
            self.text_input.tag_add("hover_token", f"1.0+{a}c", f"1.0+{b}c")

    def _on_error_select(self, event):
        self.text_input.tag_remove("error_highlight", "1.0", "end")
        sel = self.errors_tree.selection()
        if sel and self._last_error_result:
            idx = self.errors_tree.index(sel[0])
            if idx < len(self._last_error_result.errors):
                e = self._last_error_result.errors[idx]
                if e.position != (0, 0):
                    s, end = e.position
                    self.text_input.tag_add("error_highlight", f"1.0+{s}c", f"1.0+{end}c")
                    self.text_input.see(f"1.0+{s}c")

    def _set_text(self, widget, text):
        widget.configure(state="normal")
        widget.delete("1.0", "end")
        widget.insert("1.0", text)
        widget.configure(state="disabled")

    def _bind_hotkeys(self):
        self.root.bind_all("<Control-Return>", lambda _: self._run_analysis_thread())

    def clear_all(self):
        self.text_input.delete("1.0", "end")
        for r in self.tokens_table.get_children():
            self.tokens_table.delete(r)
        for r in self.errors_tree.get_children():
            self.errors_tree.delete(r)
        self._set_text(self.stats_text, "")
        self._set_text(self.report_text, "")
        self._set_text(self.internet_text, "")
        self.internet_score_label.configure(text="—")
        for lbl in self.skills_labels.values():
            lbl.configure(text="—")
        self._last_tokens, self._last_text = [], ""
        self._last_metrics, self._last_error_result = {}, None

    def show_pie_chart(self):
        if not self._last_freq:
            messagebox.showinfo("Нет данных", "Сначала выполните анализ.")
            return
        win = tk.Toplevel(self.root)
        win.title("Распределение частей речи")
        win.geometry("900x600")
        canvas = tk.Canvas(win, bg="white")
        canvas.pack(fill="both", expand=True)
        entries = [(k, float(v["коэффициент"])) for k, v in self._last_freq.items() if float(v["коэффициент"]) > 0]
        total = sum(v for _, v in entries) or 1.0
        cx, cy, r, angle = 280, 320, 200, 0.0
        for i, (_, val) in enumerate(entries):
            extent = 360.0 * (val / total)
            rgb = colorsys.hsv_to_rgb((i / max(len(entries), 1)) % 1.0, 0.65, 0.95)
            color = "#%02x%02x%02x" % tuple(int(c * 255) for c in rgb)
            canvas.create_arc(cx-r, cy-r, cx+r, cy+r, start=angle, extent=extent,
                              fill=color, outline="white", width=2)
            angle += extent
        y = 80
        canvas.create_text(560, 50, text="Легенда", anchor="w", font=("Arial", 12, "bold"))
        for i, (label, val) in enumerate(entries):
            rgb = colorsys.hsv_to_rgb((i / max(len(entries), 1)) % 1.0, 0.65, 0.95)
            color = "#%02x%02x%02x" % tuple(int(c * 255) for c in rgb)
            canvas.create_rectangle(560, y, 585, y + 18, fill=color, outline=color)
            canvas.create_text(595, y + 9, text=f"{label}: {val:.1%}", anchor="w", font=("Arial", 10))
            y += 26

    def show_heatmap(self):
        """Показать тепловую карту POS-биграмм."""
        if not self._last_metrics:
            messagebox.showinfo("Нет данных", "Сначала выполните анализ.")
            return
        pos_bg = self._last_metrics.get("pos_bigrams", {})
        if not pos_bg.get("top_bigrams"):
            messagebox.showinfo("Нет данных", "Недостаточно данных для построения heatmap.")
            return
        HeatmapWindow(self.root, pos_bg)

    def run(self):
        self.root.mainloop()


if __name__ == "__main__":
    app = AutorovedApp()
    app.run()
