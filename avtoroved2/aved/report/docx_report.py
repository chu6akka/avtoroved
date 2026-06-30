"""Заключение эксперта в формате DOCX по структуре Приложения 1 методики (с. 108–109).

Документ — вспомогательный (помощник эксперта): он систематизирует результаты стадий
и формулировку вывода; окончательное решение и подпись остаются за экспертом.
"""
from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document as Docx
from docx.shared import Pt

from aved.core.models import Level, VerdictType
from aved.core.pipeline import IdentificationResult
from aved.core.registry import Registry

_LEVEL_TITLE = {
    Level.NN: "Уровень НН (набор норм)",
    Level.NS: "Уровень НС (набор свойств норм)",
    Level.NSV: "Уровень НСВ (набор средств выражения свойств норм)",
}

_VERDICT_WORDING = {
    VerdictType.CATEGORICAL_POSITIVE:
        "Спорный текст и тексты-образцы выполнены одним и тем же лицом "
        "(категорический положительный вывод).",
    VerdictType.PROBABLE_POSITIVE:
        "Спорный текст и тексты-образцы, вероятно, выполнены одним и тем же лицом; "
        "это лицо не исключается из круга возможных авторов (вероятный положительный вывод).",
    VerdictType.PROBABLE_NEGATIVE:
        "Спорный текст и тексты-образцы, вероятно, выполнены разными лицами "
        "(вероятный отрицательный вывод).",
    VerdictType.CATEGORICAL_NEGATIVE:
        "Спорный текст и тексты-образцы выполнены разными лицами "
        "(категорический отрицательный вывод).",
    VerdictType.INCONCLUSIVE:
        "Решить вопрос об авторстве не представляется возможным "
        "(недостаточная совокупность признаков).",
}


@dataclass
class ReportMeta:
    number: str = "0000"
    city: str = "—"
    org: str = "Экспертно-криминалистическое подразделение"
    expert: str = "________________________"
    expert_credentials: str = (
        "высшее филологическое образование, экспертная специальность "
        "«Автороведческая экспертиза»"
    )
    case: str = "________________"
    basis: str = "постановление о назначении судебной автороведческой экспертизы"
    circumstances: str = (
        "Обстоятельства дела известны эксперту из постановления о назначении экспертизы."
    )
    question: str = (
        "Выполнены ли представленный спорный текст и тексты-образцы одним и тем же лицом?"
    )
    methodology: str = (
        "Рубцова И.И., Ермолова Е.И., Безрукова А.И. и др. Комплексная методика "
        "производства судебно-автороведческих экспертиз. — М.: ЭКЦ МВД России, 2007."
    )


def _feature_names(registry: Registry, ids: list[str], limit: int = 12) -> str:
    names = [registry.get(i).name for i in ids[:limit]]
    tail = f" и ещё {len(ids) - limit}" if len(ids) > limit else ""
    return ("; ".join(names) + tail) if names else "—"


def generate(
    result: IdentificationResult,
    objects,
    registry: Registry,
    path: str | Path,
    meta: ReportMeta | None = None,
) -> Path:
    meta = meta or ReportMeta()
    doc = Docx()
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(12)

    def h(text, level=1):
        doc.add_heading(text, level=level)

    def p(text=""):
        doc.add_paragraph(text)

    # --- Шапка и подписка ---
    doc.add_paragraph(meta.org).alignment = 1
    sub = doc.add_paragraph()
    sub.add_run("ПОДПИСКА").bold = True
    p(
        "Эксперту разъяснены права и обязанности, предусмотренные процессуальным "
        "законодательством. Об ответственности за дачу заведомо ложного заключения "
        "по ст. 307 УК РФ эксперт предупреждён."
    )
    p(f"Эксперт: {meta.expert}")

    title = doc.add_paragraph()
    title.alignment = 1
    title.add_run(f"ЗАКЛЮЧЕНИЕ ЭКСПЕРТА № {meta.number}").bold = True
    p(f"г. {meta.city}")

    # --- Вводная часть ---
    h("Вводная часть", 2)
    p(f"Эксперт: {meta.expert} ({meta.expert_credentials}).")
    p(f"Основание производства экспертизы: {meta.basis}.")
    p(f"Дело №: {meta.case}.")
    p(meta.circumstances)

    h("На экспертизу представлены", 2)
    for obj in objects:
        role = "спорный текст" if obj.role.value == "disputed" else "образец"
        p(f"• {obj.title} — {role}; объём {obj.word_count} слов.")

    h("Вопрос, поставленный на разрешение экспертизы", 2)
    p(meta.question)

    # --- Исследование ---
    h("ИССЛЕДОВАНИЕ", 1)
    p(
        "Исследование проведено по комплексной методике производства судебно-"
        "автороведческих экспертиз в четыре стадии: оценка пригодности объектов; "
        "раздельное исследование; сравнительное исследование по трём уровням "
        "индивидуализации навыка (НН, НС, НСВ); оценка совокупности признаков."
    )
    p(f"Методическая основа: {meta.methodology}")

    h("Стадия 1. Оценка пригодности объектов", 2)
    r = result.suitability
    p(
        f"Объём спорного материала — {r.disputed_words} слов; образцов — "
        f"{r.sample_words} слов; соотношение ×{r.volume_ratio}."
    )
    for note in r.notes:
        p(f"• {note}")
    if not r.notes:
        p("Объекты признаны пригодными для идентификационного исследования.")

    if result.comparison is None or result.verdict is None:
        h("ВЫВОДЫ", 1)
        p("Исследование не доведено до сравнительной стадии (объекты непригодны).")
        out = Path(path)
        doc.save(out)
        return out

    h("Стадия 3. Сравнительное исследование", 2)
    p(
        "Сопоставление моделей навыка спорного текста и образцов по уровням "
        "индивидуализации. Совпадающие и различающиеся признаки:"
    )
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Уровень"
    hdr[1].text = "Совпадения"
    hdr[2].text = "Различия"
    hdr[3].text = "Высокоинф. совпадения"
    for lv in (Level.NN, Level.NS, Level.NSV):
        lc = result.comparison.levels[lv]
        row = table.add_row().cells
        row[0].text = _LEVEL_TITLE[lv]
        row[1].text = str(len(lc.matching))
        row[2].text = str(len(lc.differing))
        row[3].text = str(lc.matching_high)

    if result.comparison.nn_norm_conflict:
        p(f"Конфликт на уровне НН: {result.comparison.nn_conflict_reason}.")

    for lv in (Level.NN, Level.NS, Level.NSV):
        lc = result.comparison.levels[lv]
        h(_LEVEL_TITLE[lv], 3)
        p("Совпадающие признаки: " + _feature_names(registry, lc.matching))
        p("Различающиеся признаки: " + _feature_names(registry, lc.differing))

    # --- Выводы ---
    h("ВЫВОДЫ", 1)
    v = result.verdict
    concl = doc.add_paragraph()
    concl.add_run(_VERDICT_WORDING[v.type]).bold = True
    p(
        f"Высокоинформативных совпадений: {v.matching_high_count} "
        f"(порог методики — {v.MIN_HIGH_INFO}). "
        f"Высокоинформативных различий: {v.differing_high_count}."
    )
    h("Обоснование", 2)
    for line in v.rationale:
        p(f"• {line}")

    p()
    p(
        "Выводы носят вспомогательный характер; окончательная оценка совокупности "
        "признаков и формулирование вывода относятся к компетенции эксперта."
    )
    p(f"Эксперт: {meta.expert} / подпись ______________")

    out = Path(path)
    doc.save(out)
    return out
