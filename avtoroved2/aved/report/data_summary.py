"""Сводка лингвистических показателей (сухие данные, без выводов) в DOCX.

Это НЕ заключение и НЕ черновик вывода: только верифицируемые цифры и найденные
маркеры с примерами. Оценку и выводы делает эксперт.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document as Docx
from docx.shared import Pt


def generate(profiles: dict, path: str | Path) -> Path:
    doc = Docx()
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(11)

    doc.add_heading("Сводка лингвистических показателей", level=1)
    doc.add_paragraph(
        "Справочные данные для эксперта. Программа не делает выводов об авторстве; "
        "оценка совокупности признаков и выводы — компетенция эксперта."
    )

    ids = list(profiles)
    if not ids:
        out = Path(path)
        doc.save(out)
        return out

    doc.add_heading("Объекты", level=2)
    for pid in ids:
        p = profiles[pid]
        role = "спорный" if p["role"] == "disputed" else "образец"
        doc.add_paragraph(f"• {p['title']} — {role}; {p['word_count']} слов")

    # Таблица показателей: строки — показатели, столбцы — объекты
    doc.add_heading("Количественные показатели", level=2)
    metric_names = list(profiles[ids[0]]["metrics"])
    table = doc.add_table(rows=1, cols=1 + len(ids))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Показатель"
    for j, pid in enumerate(ids):
        hdr[j + 1].text = profiles[pid]["title"]
    for name in metric_names:
        row = table.add_row().cells
        row[0].text = name
        for j, pid in enumerate(ids):
            row[j + 1].text = str(profiles[pid]["metrics"].get(name, ""))

    # Маркеры (только найденные, с примерами)
    doc.add_heading("Лексические маркеры (фактически найденные)", level=2)
    for pid in ids:
        p = profiles[pid]
        doc.add_heading(p["title"], level=3)
        if not p["markers"]:
            doc.add_paragraph("— маркеры не обнаружены")
            continue
        for m in p["markers"]:
            doc.add_paragraph(
                f"• {m['name']}: {m['count']} (на 1000 слов: {m['rate']}; {m['source']}). "
                f"Примеры: {', '.join(m['examples'])}"
            )

    out = Path(path)
    doc.save(out)
    return out
